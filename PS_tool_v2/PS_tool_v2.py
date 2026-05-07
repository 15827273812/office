import tkinter as tk
from tkinter import messagebox
import tkinter.ttk as ttk
from tkinter import filedialog
from venv import logger
import pandas as pd
from openpyxl import Workbook
import pyodbc
import re
import os
from datetime import datetime as dt
import subprocess
import shutil
import pyautogui
import time
from PIL import Image
import tempfile
import pyperclip
import datetime
import easygui
from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver.support.ui import WebDriverWait, Select
from selenium.webdriver.edge.service import Service
from selenium.webdriver.common.keys import Keys
from webdriver_manager.microsoft import EdgeChromiumDriverManager
from PIL import Image
import glob
import docx
import xlsxwriter
import requests
import json
import pathlib
import sys

# ============================
# 从 config.json 加载配置
# ============================
SCRIPT_DIR = pathlib.Path(__file__).parent.absolute()
CONFIG_PATH = SCRIPT_DIR / 'config.json'

if not CONFIG_PATH.exists():
    print(f"错误: 找不到配置文件 {CONFIG_PATH}")
    print("请确保 config.json 与脚本在同一目录下")
    sys.exit(1)

with open(CONFIG_PATH, 'r', encoding='utf-8') as _f:
    CONFIG = json.load(_f)

CON_STRING = CONFIG['database']['connection_string']
SQL_DATA = CONFIG['sql_queries']
SN = CONFIG['service_now']
VC = CONFIG['voice_console']
NIKE_EPS = CONFIG['nike_eps']
HC = CONFIG['health_check']

con_string = CON_STRING  # 向后兼容




window = tk.Tk()
window.title("工作台")
# 获取屏幕尺寸
screen_width = window.winfo_screenwidth()
screen_height = window.winfo_screenheight()

# 计算窗口尺寸
window_width = 1500
window_height = 1090

# 计算窗口位置
pos_right = int(screen_width / 2 - window_width / 2)
pos_down = int(screen_height / 2 - window_height / 2)

# 设置窗口位置属性
window.geometry(f"{window_width}x{window_height}+{pos_right}+{pos_down}")

style = ttk.Style()

# ('winnative', 'clam', 'alt', 'default', 'classic', 'vista', 'xpnative')

style.theme_use('xpnative')

style.configure('TNotebook', background='lightgray')

style.map('TNotebook.Tab', background=[('selected', 'lightgreen')], foreground=[('selected', 'darkblue')])

########################################################
#  notebook area
########################################################


notebook        = ttk.Notebook(window)

data_export     = ttk.Frame(notebook)

#ping_ip         = ttk.Frame(notebook)

label_history   = ttk.Frame(notebook)



########################################################

#  function area
########################################################

def Show_Page(event):

    selected_tab = notebook.tab(notebook.select(), "text")

    # print(f"Selected tab: {selected_tab}")


# def Open_File():

#     file_path = filedialog.askopenfilename(

#         filetypes=[

#             ("Text files", "*.txt"),

#             ("Sql files", "*.sql"),

#             ("All files", "*.*")

#             ]
#     )

#     if not file_path:

#         window.title("无标题 - 工作台")

#         return
    
#     sql.delete("1.0",tk.END)

#     with open(

#         file_path,
#         mode="r",

#         encoding="utf-8") as input_file:

#         sql.insert(

#             tk.END,

#             input_file.read()
#             )

#     message = f'{dt.now()}: 已打开 {file_path} !\n'
#     log.insert("insert", message)



# def Copy():

#     selected_text = sql.get("sel.first", "sel.last")

#     clipboard = window.clipboard_append(selected_text)

#     message = f'{dt.now()}: 已复制 !\n'
#     log.insert("insert", message)

#     # messagebox.showinfo("提示", "已复制到剪贴板")


# def Paste():

#     clipboard_text = window.clipboard_get()

#     sql.insert("insert", clipboard_text)

#     message = f'{dt.now()}: 已粘贴 !\n'
#     log.insert("insert", message)



def Clear():

    sql.delete("1.0",tk.END)

    message = f'{dt.now()}: 清屏完成 !\n'
    log.insert("insert", message)


def Find_Path():

    file_path = filedialog.askdirectory()

    path.delete(0, "end")

    path.insert(0, file_path)

def bytes_to_string(byte_val):
    if isinstance(byte_val, bytes):  # 仅当值为bytes类型时才解码
        try:
            return byte_val.decode('cp037')
        except UnicodeDecodeError:
            return ''
    elif isinstance(byte_val, str):  # 如果已经是字符串，直接返回
        return byte_val
    else:  # 其他类型（如int、float），转换为字符串
        return str(byte_val)

def process_data_with_sql(df):
    try:
        numeric_columns = ['INONHQTY', 'DU', 'PU', 'QTY','FREEQTY','BLOCKQTY','BPPIDQTY','INUPCNBR','IBIEXQTY','ASIACQTY','SHACPQTY','UNITQTY']
    # 进行数字类型转换
        for column in numeric_columns:
            if column in df.columns:
                df[column] = pd.to_numeric(df[column], errors='coerce')
    except Exception as e:
        logger.error(f"Error during data processing: {e}")
        # 返回原DataFrame，或一个空的DataFrame
        return pd.DataFrame()
    return df


def Formate():

    def add_prefix(match_obj):

        return 'ctxprwdta.' + match_obj.group(0)

    input_str = sql.get('1.0', tk.END)

    pattern = r'\b(?<![ctxprwdta.])\w{2}[rRwW]\w{3}\b'

    result = re.sub(pattern, add_prefix, input_str)

    sql.delete("1.0",tk.END)

    sql.insert("insert", result)

    message = f'{dt.now()}: 格式化完成 !\n'
    log.insert("insert", message)

def GI_status():
    log.delete(1.0, tk.END)
    try:
        # 连接到数据库
        cnxn = pyodbc.connect(con_string)
        cur = cnxn.cursor()
        
        # 第一个 SQL 查询
        sql_query_1 = SQL_DATA['gi_status']['summary']
        
        cur.execute(sql_query_1)
        rows = cur.fetchall()
        
        # 新增的 SQL 查询，获取还没有进行 GI 的 PL 数量
        sql_query_6 = SQL_DATA['gi_status']['gi_status']
        cur.execute(sql_query_6)
        pl_status_rows = cur.fetchall()
        # 检查第一列的值是否只包含 '160' 和 '090'
        gnstscde_values = set(row[0] for row in rows)
        pl_status_values = set(row[0] for row in pl_status_rows)
        if gnstscde_values.issubset({'160', '090'}) and pl_status_values.issubset({'Y'}):
            # 检查 GNSERR 表中是否有 GNERRLNK 为 'GIPROCESS' 的记录
            sql_query_5 = SQL_DATA['gi_status']['errors']
            cur.execute(sql_query_5)
            error_rows = cur.fetchall()
            
            if not error_rows:
                # 执行后续的 SQL 查询
                sql_query_2 = SQL_DATA['gi_status']['max_cmr_time']
                cur.execute(sql_query_2)
                max_shccmtme = cur.fetchone()[0]
                
                sql_query_3 = SQL_DATA['gi_status']['today_volume']
                cur.execute(sql_query_3)
                sum_shacpqty, count_inhldcde, count_dndnhnbr, count_plpklnbr = cur.fetchone()
                
                sql_query_4 = SQL_DATA['gi_status']['max_gi_time']
                cur.execute(sql_query_4)
                max_gnjobtme = cur.fetchone()[0]
                if max_shccmtme<max_gnjobtme:
                    # 打印固定的内容
                    message = (
                    "Hi all\n"
                        "\n"
                        "Please find below EOD status:\n"
                        "1. All GI completed.\n"
                        "2. GI completed data successfully sent out from WMS\n"
                        "3. Last CMR at {}\n"
                        "4. Last GI at {}\n"
                        "5. Last GI volume for today: {} units, {} cartons, {} DNs, {} PLs.\n"
                    ).format(f"{str(max_shccmtme)[:2]}:{str(max_shccmtme)[2:4]}:{str(max_shccmtme)[4:6]}", f"{str(max_gnjobtme)[:2]}:{str(max_gnjobtme)[2:4]}:{str(max_gnjobtme)[4:6]}", sum_shacpqty, count_inhldcde, count_dndnhnbr, count_plpklnbr)
                    log.insert("insert", message)
                else:
                    message = ("GI未结束:CMR时间大于GI时间！")
                    log.insert("insert", message)
            else:
                message = ("GI未结束,存在GI ERROR")
                log.insert("insert", message)
        else:
            message = ("GI未结束:\n")
            log.insert("insert", message)
            gi_not_completed_message = "STSCODE    FLOW TYPE       COUNT\n"
            total_cartons = 0
            for row in rows:
                STSCODE, FLOW_TYPE, CARTONS = row
                gi_not_completed_message += f"{STSCODE:>5}   {FLOW_TYPE:>11} {CARTONS:>11}\n"
                if STSCODE != '090':
                    total_cartons += CARTONS
            gi_not_completed_message += f"{'total:':>5} {total_cartons:>4}\n"
            log.insert("insert", gi_not_completed_message)
            # 构建未进行 GI 的 PL 数量信息
            pl_status_message = "\nGI Status:\n"
            pl_status_message += "SHGISFLG    COUNT\n"
            for row in pl_status_rows:
                SHGISFLG, count = row
                pl_status_message += f"{SHGISFLG:>5}   {count:>8}\n"
            
            log.insert("insert", pl_status_message)
        
        # 关闭数据库连接
        cur.close()
        cnxn.close()
    
    except Exception as e:
        message = (f'Error: {str(e)}')
        log.insert("insert", message)

def create_pictures_folder():
    # 创建 pictures 文件夹（如果不存在）
    if not os.path.exists('pictures'):
        os.makedirs('pictures')
    else:
        files = glob.glob('pictures/*')
        for f in files:
            if os.path.isfile(f):
                os.remove(f)

def copy_page_text():
    # 全选并复制当前页面的文字
    pyautogui.hotkey('ctrl', 'a')
    pyautogui.hotkey('ctrl', 'c')
    # 获取剪贴板上的文本
    clipboard_text = pyperclip.paste()
    # 将文本写入临时文件
    with tempfile.NamedTemporaryFile(mode='w+', delete=False,encoding="utf - 8") as temp_file:
        temp_file.write(clipboard_text)
        temp_file.flush()
        # 返回临时文件的名称
        return temp_file.name
def add_screenshot_to_word(screenshot_path, doc):
    # 将图片添加到 Word 文档中
    doc.add_picture(screenshot_path, width=docx.shared.Inches(4.9),height =docx.shared.Inches(3.0))
    
def create_screenshots(job_name,page,doc=None): 
        global should_copy_text           
        # 初始化 job 的起始位置和结束位置
        job_start_y = 100  # 假设 job 起始于屏幕顶部
        job_height = 980   # 假设每个 job 高度为 100 像素
        sheet_name = job_name.replace(' ', '_')[:31]  # Excel 工作表名不能超过 31 个字符     
        # 复制当前页面的文字
        if should_copy_text:
            temp_file_name = copy_page_text()
            with open(temp_file_name, 'r',encoding="utf - 8") as temp_file:
                 temp_file.read()
            # 删除临时文件
            os.remove(temp_file_name)
        # 截取 job 区域并调整大小后保存到 Excel
        screenshot = pyautogui.screenshot(region=(0, job_start_y, pyautogui.size().width, job_height))
        # 调整图片大小
        screenshot = screenshot.resize((screenshot.width // 2, screenshot.height // 2), Image.Resampling.LANCZOS)
        screenshot_path = os.path.join('pictures', f'screenshot_{sheet_name}_{page}.png')
        # 保存图片时设置质量
        screenshot.save(screenshot_path, optimize=True, quality=60)  # 优化并设置质量
        # 将图片添加到 Word 文档中
        if doc is not None:
            add_screenshot_to_word(screenshot_path, doc)
        time.sleep(1)


def take_screenshots(jobs):
    for job_name in jobs:
        page = 1

        if job_name=='STRSQL':
            #输入命令
            pyautogui.press('tab')
            pyautogui.typewrite(job_name)
            pyautogui.press('enter')

            #输入F14=Confirm
            pyautogui.keyDown('shift')
            pyautogui.press('f2')
            pyautogui.press('tab')
            pyautogui.keyUp('shift')
            #截图并退到命令行界面
            doc.add_paragraph(f"1.check if system can login or not (both job role menu and STRSQL menu)两张")
            time.sleep(0.5)
            create_screenshots(job_name,page,doc)
            pyautogui.press('f12')
            pyautogui.press('enter')

        if job_name=='XPDS CTXPR':
            #输入命令并进到JOBROLE界面
            pyautogui.press('tab')
            pyautogui.typewrite(job_name)
            pyautogui.press('enter')
            pyautogui.press('enter')

            #截图并退到命令行界面
            time.sleep(0.5)
            create_screenshots(job_name,page,doc)
            pyautogui.press('f3')

        if job_name=='WRKMQM':
            #输入命令
            pyautogui.press('tab')
            pyautogui.typewrite(job_name)
            pyautogui.press('enter')

            #截图并退到命令行界面
            doc.add_paragraph(f"4.check MQ 第一页(WRKMQM)")
            time.sleep(0.5)
            create_screenshots(job_name,page,doc)
            pyautogui.press('f3')

        if job_name=='WRKACTJOB SBS(CTXPRCDC)':
            #输入命令
            pyautogui.press('tab')
            pyautogui.typewrite(job_name)
            pyautogui.press('enter')
            pyautogui.press('f7')
            pyautogui.typewrite('XPDSN')
            pyautogui.press('enter')
            time.sleep(0.5)
            #截图并退到命令行界面
            doc.add_paragraph(f"WRKACTJOB SBS(CTXPRCDC):")
            time.sleep(0.5)
            create_screenshots(job_name,page,doc)
            pyautogui.press('f3')

        if job_name=='DSPFD FILE(SVSDLNA)':
            #输入命令
            pyautogui.press('tab')
            pyautogui.typewrite(job_name)
            pyautogui.press('enter')

            #输入B到最后一页
            pyautogui.typewrite("B")
            pyautogui.press('enter')
            time.sleep(1)

            #截图并退到命令行界面
            doc.add_paragraph(f"7.DSPFD FILE(SVSDLNA) 最后一页")
            time.sleep(0.5)
            create_screenshots(job_name,page,doc)
            pyautogui.press('f3')

        if job_name=='DSPMSG MSGQ(*SYSOPR) SEV(99)':
            #输入命令            
            pyautogui.press('tab')
            pyautogui.typewrite(job_name)
            pyautogui.press('enter')
            time.sleep(2)

            #截图并退到命令行界面
            doc.add_paragraph(f"9.DSPMSG MSGQ(*SYSOPR) SEV(99):")
            time.sleep(0.5)
            create_screenshots("SYSOPR",page,doc)
            pyautogui.press('f3')

        if job_name=='DSPMSG MSGQ(CTXPRWSCD) SEV(99)':
            #输入命令
            pyautogui.press('tab')
            pyautogui.typewrite(job_name)
            pyautogui.press('enter')
            time.sleep(5)

            #截图并退到命令行界面
            doc.add_paragraph(f"DSPMSG MSGQ(CTXPRWSCD) SEV(99)")
            time.sleep(0.5)
            create_screenshots("CTXPRWSCD",page,doc)
            pyautogui.press('f3')

        if job_name=='WRKLNK':
            #输入命令WRKLNK
            pyautogui.press('tab')
            pyautogui.typewrite(job_name)
            pyautogui.press('enter')
            time.sleep(0.5)

            #查找home行并输入5=display
            for _ in range(7):
                pyautogui.press('down')
            pyautogui.typewrite("5")
            pyautogui.press('enter')
            time.sleep(0.5)

            #查找CHINA行并输入5=display
            pyautogui.keyDown('ctrl')
            pyautogui.press('q')
            pyautogui.keyUp('ctrl')
            pyautogui.press('down')    
            pyautogui.typewrite("5")
            pyautogui.press('enter')
            time.sleep(0.5)

            #查找PRD行并输入5=display
            pyautogui.press('down')    
            pyautogui.press('down')    
            pyautogui.typewrite("5")
            pyautogui.press('enter')
            time.sleep(0.5)

            #截图并退到命令行界面
            doc.add_paragraph(f"8.WRKLNK/home/CHINA/PRD")
            time.sleep(0.5)
            create_screenshots(job_name,page,doc)
            pyautogui.press('f3')
            pyautogui.press('enter')

        if job_name=='General Link':
            #输入命令进入到JOBROLE界面
            pyautogui.press('tab')
            pyautogui.typewrite("XPDS CTXPR")
            pyautogui.press('enter')
            pyautogui.press('enter')

            #进入LINK界面
            pyautogui.typewrite("60")
            pyautogui.press('enter')
            pyautogui.typewrite("10")
            pyautogui.press('enter')
            pyautogui.typewrite("60")
            pyautogui.press('enter')
            pyautogui.typewrite("50")
            pyautogui.press('enter')
            pyautogui.typewrite("10")
            pyautogui.press('enter')
            #STATUS行输入F
            for _ in range(3):
                pyautogui.press('tab')
            pyautogui.typewrite("F")
            time.sleep(0.5)
            #清空日期和时间并进入结果画面
            pyautogui.press('tab')
            pyautogui.press('tab')
            for _ in range(16):
                pyautogui.press('space')
            pyautogui.press('enter')

            #截图并退到命令行界面
            doc.add_paragraph(f"3.check general link have F status or not:")
            time.sleep(0.5)
            create_screenshots(job_name,page,doc)
            pyautogui.press('f3')
            pyautogui.press('f3')
            pyautogui.press('enter')

        if job_name=='WRKACTJOB':

            #输入命令进入到一览界面
            pyautogui.press('tab')
            pyautogui.typewrite(job_name)
            pyautogui.press('enter')

            #到达STATUS列并排序
            for _ in range(63):
                pyautogui.press('right')
            pyautogui.keyDown('shift')
            pyautogui.press('f4')
            pyautogui.keyUp('shift')

            #查找ROBOTREACT
            pyautogui.press('f7')
            pyautogui.typewrite("MSGW")
            pyautogui.press('tab')
            pyautogui.typewrite("*STS")
            pyautogui.press('enter')
            paragraph = doc.add_paragraph()
            paragraph.add_run(f"6.wrkactjob check MSGW job\nWRKACTJOB:")
            time.sleep(0.5)
            create_screenshots(job_name,page,doc)
            page += 1

            pyautogui.press('f3')

        if job_name=='WATCHDOG':
            #输入命令进入JOBROLE界面
            pyautogui.press('tab')
            pyautogui.typewrite("XPDS CTXPR")
            pyautogui.press('enter')
            pyautogui.press('enter')

            #进入WATCHDOG界面
            pyautogui.typewrite("60")
            pyautogui.press('enter')
            pyautogui.typewrite("110")
            pyautogui.press('enter')
            pyautogui.typewrite("20")
            pyautogui.press('enter')
            pyautogui.typewrite("50")
            pyautogui.press('enter')

            #对象DOG输入13
            pyautogui.press('tab')
            pyautogui.typewrite("13")
            pyautogui.press('tab')
            pyautogui.press('tab')
            for _ in range(6):
                pyautogui.typewrite("13")
            pyautogui.press('enter')

            #截图
            page_down_count = 0
            f12_count = 1
            first_screenshot_after_f12 = True
            for page_down_count in range(0,14):
                #第一页截图
                if page_down_count == 0:
                   doc.add_paragraph(f"2.check WD CTXPRILM:两张")
                   time.sleep(0.5)
                   create_screenshots(job_name,page,doc)
                   page += 1

                #判读是否能翻页
                if page_down_count <=14:
                    pyautogui.keyDown('ctrl')
                    pyautogui.press('q')
                    pyautogui.keyUp('ctrl')
                    page_down_count += 1
                    
                #翻页后如果是截图界面则截图
                if page_down_count in [1 , 3 , 7 , 8 , 13 , 14] :
                    time.sleep(0.5)
                    create_screenshots(job_name,page,doc)
                    page += 1

                #判断翻页后是否可以按F12跳转到下一个WATCHDOG
                if page_down_count in [1 , 3 , 8 , 13 , 14] :
                    pyautogui.press('f12')
                    if first_screenshot_after_f12:
                        # 根据 f12_count 添加不同的文字
                        if f12_count == 1:
                            doc.add_paragraph(f"CTXPRIML:2张")
                        elif f12_count == 2:
                            doc.add_paragraph(f"CTXPRIMM:3张,第一和最后两张")
                        elif f12_count == 3:
                            doc.add_paragraph(f"CTXPRIMP:2张(第一和最后一张)")
                        elif f12_count == 4:
                            doc.add_paragraph(f"CTXPRIMS:2张")
                        elif f12_count == 5:
                            doc.add_paragraph(f"CTXPRIMV:第一页")
                        first_screenshot_after_f12 = False
                    #F12跳转后截图并判断是否是最后一个WATCHDOG
                    if (f12_count <= 5):
                        first_screenshot_after_f12 = True
                        time.sleep(0.5)
                        create_screenshots(job_name,page,doc)
                        page += 1
                        f12_count += 1

                    if (f12_count == 6):
                        pyautogui.press('f12')
                        doc.add_paragraph(f"CTXPRWMO 1张") 
                        time.sleep(0.5)
                        create_screenshots(job_name,page,doc)

            pyautogui.press('f3')
            pyautogui.press('f3')

        if job_name=='RBM':

            #进入RBM界面
            pyautogui.typewrite(job_name)
            pyautogui.press('enter')
            pyautogui.typewrite("1")
            pyautogui.press('enter')
            time.sleep(0.5)

            today = datetime.datetime.now()
            weekday_num = today.weekday()

            # #如果当天是周一，则对#WSUN截图
            # if weekday_num in [0]:
            #    pyautogui.typewrite("#WSUN")
            #    pyautogui.press('enter')
            #    pyautogui.typewrite("11")
            #    pyautogui.press('enter')
            #    time.sleep(0.5)
            #    paragraph = doc.add_paragraph()
            #    paragraph.add_run(f"5.check daily/weekly/sunday/monthly/CDC job is complete or not\n#WSUN:")
            #    time.sleep(0.5)
            #    create_screenshots(job_name,page,doc)
            #    page += 1
            #    pyautogui.press('f3')

            #如果当天是周二到周六，则对#DAILYPR截图
            if weekday_num in [0,1,2,3,4,5,6]:
               pyautogui.typewrite("#DAILYPR")
               pyautogui.press('enter')
               pyautogui.typewrite("11")
               pyautogui.press('enter')
               time.sleep(0.5)
               paragraph = doc.add_paragraph()
               paragraph.add_run(f"5.check daily/weekly/sunday/monthly/CDC job is complete or not\n#DAILYPR:")
               time.sleep(0.5)
               create_screenshots(job_name,page,doc)
               page += 1
               pyautogui.press('f3')

            # #如果当天是周日，则对#WKLY2~5截图
            # if weekday_num in [6]:
            #    for var_name in range(2,6):
            #        pyautogui.typewrite(f"#WKLY{var_name}")
            #        pyautogui.press('enter')
            #        time.sleep(0.5)

            #        pyautogui.typewrite("11")
            #        pyautogui.press('enter')
            #        time.sleep(0.5)
            #        paragraph = doc.add_paragraph()
            #        paragraph.add_run(f"5.check daily/weekly/sunday/monthly/CDC job is complete or not\n#WKLY{var_name}")
            #        time.sleep(0.5)
            #        create_screenshots(job_name,page,doc)
            #        page += 1
            #        pyautogui.press('f3')

            #        pyautogui.press('f3')
            #        pyautogui.typewrite("1")
            #        pyautogui.press('enter')
            #        time.sleep(0.5)

            #CTXPRG_STR截图
            pyautogui.press('f3')
            pyautogui.typewrite("1")
            pyautogui.press('enter')
            pyautogui.typewrite("CTXPRG_STR")
            pyautogui.press('enter')
            pyautogui.typewrite("11")
            pyautogui.press('enter')
            time.sleep(0.5)
            doc.add_paragraph(f"CTXPRG_STR:")
            time.sleep(0.5) 
            create_screenshots(job_name,page,doc)
            pyautogui.press('f3')

            pyautogui.press('f3')
            pyautogui.press('f3')
    
def open_and_screenshot():
    # 设置浏览器选项 - 提高性能
      options = webdriver.ChromeOptions()
      options.add_experimental_option("prefs", {
      # 禁用密码保存提示
      "credentials_enable_service": False,
      "profile.password_manager_enabled": False,
      })
      options.add_argument('--disable-blink-features=AutomationControlled')
      options.add_experimental_option("excludeSwitches", ["enable-automation"])
      options.add_experimental_option('useAutomationExtension', False)
      options.add_argument('--disable-extensions')
      options.add_argument('--no-sandbox')
      options.add_argument('--disable-dev-shm-usage')
      options.add_argument('--disable-gpu')

      # 设置浏览器驱动
      driver_path=None
      if driver_path:
           driver = webdriver.Chrome(executable_path=driver_path, options=options)
      else:
           driver = webdriver.Chrome(options=options)

  # 隐藏自动化特征
      driver.execute_script("Object.defineProperty(navigator, 'webdriver', {get: () => undefined})")

      global should_copy_text
      should_copy_text = False
      try:
          # 打开指定网址
          driver.maximize_window()
          driver.get("http://tcgiapp1wapp013:9090/VoiceConsole/login.action")
          # 使用显式等待替代固定等待
          wait = WebDriverWait(driver, 10)
          # 定位登录表单元素
          username = wait.until(EC.presence_of_element_located((By.NAME, "j_username")))  # 根据实际元素属性修改
          password = wait.until(EC.presence_of_element_located((By.NAME, "j_password")))  # 根据实际元素属性修改
        
          # 直接使用Selenium输入
          username.send_keys(VC["username"])
          password.send_keys(VC["password"])
          password.submit()  # 自动提交表单
          # 等待页面加载
          time.sleep(3)
          # 获取页面截图
          paragraph = doc.add_paragraph()
          paragraph.add_run(f"10.check Voice web\nhttp://tcgiapp1wapp013:9090/VoiceConsole/login.action")
          time.sleep(0.5)
          create_screenshots("Honeywell",1,doc)

          driver.get("https://techops-clcps.nike.com/#/dashboard")
          time.sleep(1)
           # 定位新网站的登录元素
          nike_user = wait.until(EC.presence_of_element_located((By.XPATH, "//input[@class='el-input__inner' and @placeholder='用户名 / 手机 / 邮箱']")))

        # 根据实际元素属性修改
          nike_pass = wait.until(EC.presence_of_element_located((By.XPATH, "//input[@class='el-input__inner' and @placeholder='请输入密码']")))
  # 根据实际元素属性修改
          nike_user.send_keys(NIKE_EPS["username"])
          nike_pass.send_keys(NIKE_EPS["password"])
          wait.until(EC.element_to_be_clickable((By.XPATH, "//span[text()='登录']/ancestor::button[contains(@class,'el-button--primary')]"))).click()
          time.sleep(3)
          # 获取页面截图
          paragraph = doc.add_paragraph()
          paragraph.add_run(f"11.check EPS server(first lic server TCGWMSP1WEPS003)\n远程用户/密码\nEe*GC7qz\n\nTCGWMSP1WEPS003\n\n12.https://techops-clcps.nike.com/#/dashboard")
          time.sleep(0.5)
          create_screenshots("dashboard",1,doc)

          x, y = 900, 700
          pyautogui.moveTo(x, y, duration=3)
          time.sleep(0.5)
          pyautogui.click()
          time.sleep(2)
          # 获取页面截图
          time.sleep(0.5)
          create_screenshots("dashboard",2,doc)
      except Exception:
          print("we can't open this path")
      finally:
        # 关闭浏览器
        driver.close()

      try:
          path = HC["log_path"]
          os.startfile(path)
          time.sleep(5)
          pyautogui.hotkey('ctrl','end')
          time.sleep(3)
          paragraph = doc.add_paragraph()
          paragraph.add_run(f"13.\\tcpsappd1ap01\deploy\ backend\log-app.log\n查看这个log文件内容是否是最近5分钟之内更新的")
          time.sleep(0.5)
          create_screenshots("log",1,doc)
          time.sleep(3)
          pyautogui.hotkey('ctrl','W')
      except Exception:
          print("we can't open this path")
      easygui.msgbox("截图结束！请检查结果！！！")
    
should_copy_text = True
current_date = dt.now().strftime("%m%d")
word_doc_path = f"{current_date}.docx"
doc = docx.Document()


def health_check():
    choice = easygui.buttonbox(
        msg="请切换到AS400 Command Entry界面并确认翻页快捷键已设置!!",
        title="确认",
        choices=["OK", "Cancel"]
    )
    
    # 如果用户点击了 Cancel 或关闭了弹窗，则终止函数
    if choice == "Cancel" or choice is None:
        return

    time.sleep(2)
    jobs = [
        "STRSQL",
        "XPDS CTXPR",
        "WATCHDOG",
        "General Link",
        "WRKMQM",
        "RBM",
        "WRKACTJOB SBS(CTXPRCDC)",
        "WRKACTJOB",
        "DSPFD FILE(SVSDLNA)",
        "WRKLNK",
        "DSPMSG MSGQ(*SYSOPR) SEV(99)",
        "DSPMSG MSGQ(CTXPRWSCD) SEV(99)"
    ]
    # AS400截图
    # 创建pictures文件夹
    create_pictures_folder()  
    take_screenshots(jobs)
    # 网页截图
    open_and_screenshot()  
    # 关闭DOC
    doc.save(word_doc_path)
    shutil.rmtree('pictures')

class ExcelDeleteAutomation:
    def __init__(self,driver_path=None):
        self.data = self.read_data()
        
        # 设置浏览器选项 - 提高性能
        options = webdriver.ChromeOptions()
        options.add_experimental_option("prefs", {
            # 禁用密码保存提示
            "credentials_enable_service": False,
            "profile.password_manager_enabled": False,
        })
        options.add_argument('--disable-blink-features=AutomationControlled')
        options.add_experimental_option("excludeSwitches", ["enable-automation"])
        options.add_experimental_option('useAutomationExtension', False)
        options.add_argument('--disable-extensions')
        options.add_argument('--no-sandbox')
        options.add_argument('--disable-dev-shm-usage')
        options.add_argument('--disable-gpu')
        
        # 设置浏览器驱动
        if driver_path:
            self.driver = webdriver.Chrome(executable_path=driver_path, options=options)
        else:
            self.driver = webdriver.Chrome(options=options)
        
        # 隐藏自动化特征
        self.driver.execute_script("Object.defineProperty(navigator, 'webdriver', {get: () => undefined})")
        
        # 使用更短的等待时间
        self.wait = WebDriverWait(self.driver, 5)  # 从10秒减少到5秒
        self.short_wait = WebDriverWait(self.driver, 2)  # 短等待
    
    def read_data(self):
        try:
            input_values = [c.strip() for c in sql.get("1.0", tk.END).split('\n') if c.strip()]
            return input_values
        except Exception as e:
            message = (f"读取data失败: {e}")
            log.insert("insert", message)
            return []
    
    def login(self):
        try:
          # 打开指定网址
          self.driver.get("http://tcgiapp1wapp013:9090/VoiceConsole/login.action")
          self.driver.maximize_window()
          # 使用显式等待替代固定等待
          wait = WebDriverWait(self.driver, 10)
          # 定位登录表单元素
          username = wait.until(EC.presence_of_element_located((By.NAME, "j_username")))  # 根据实际元素属性修改
          password = wait.until(EC.presence_of_element_located((By.NAME, "j_password")))  # 根据实际元素属性修改
        
          # 直接使用Selenium输入
          username.send_keys(VC["username"])
          password.send_keys(VC["password"])
          password.submit()  # 自动提交表单
          # 等待页面加载
          time.sleep(1)
          return True     
        except Exception as e:
            message = (f"登录过程中出错: {e}")
            log.insert("insert", message)
            return False

    
    def check_no_results(self):
        """
        快速检查是否显示"没有得到任何结果"
        """
        try:
            # 快速检查页面文本
            body_text = self.driver.find_element(By.TAG_NAME, "body").text
            no_result_indicators = ["没有得到任何结果", "未找到任何结果", "没有找到数据", "查询无结果"]
            
            for indicator in no_result_indicators:
                if indicator in body_text:
                    return True
            return False
        except:
            return False
    
    def find_and_paste_text(self, text, input_selector="input[type='text']"):
        """
        快速查找输入框并粘贴文本
        """
        try:
            # 快速查找输入框
            input_selectors = [
                "input[type='text']",
                "input[type='search']",
                "input[placeholder*='搜索']",
                "input[placeholder*='查询']",
                ".search-input",
                "#search"
            ]
            
            input_element = None
            for selector in input_selectors:
                try:
                    input_element = self.driver.find_element(By.CSS_SELECTOR, selector)
                    if input_element.is_displayed():
                        break
                except:
                    continue
            
            if not input_element:
                # 尝试查找第一个输入框
                try:
                    input_element = self.driver.find_element(By.TAG_NAME, "input")
                except:
                    return False
            
            # 快速输入（使用直接输入而不是复制粘贴）
            input_element.clear()
            input_element.send_keys(str(text))
            input_element.send_keys(Keys.RETURN)
            time.sleep(1)  # 减少等待时间
            
            # 快速检查无结果
            if self.check_no_results():
                return "no_results"
            
            return True
        except Exception as e:
            return False
    
    def fast_click_by_text(self, text):
        """
        快速点击包含指定文本的元素
        """
        try:
            # 同时查找多种可能包含文本的元素
            element_types = ['a', 'span', 'div', 'td', 'button', 'li']
            
            for elem_type in element_types:
                try:
                    element = self.driver.find_element(By.XPATH, f"//{elem_type}[contains(text(), '{text}')]")
                    if element.is_displayed():
                        element.click()
                        time.sleep(0.5)  # 减少等待时间
                        return True
                except:
                    continue
            return False
        except Exception as e:
            return False
    
    def fast_check_and_delete(self):
        """
        快速检查并执行删除操作
        """
        try:
            # 快速检查页面文本
            page_text = self.driver.find_element(By.TAG_NAME, "body").text
            
            if "未分配设备" in page_text:
                start_time = time.time()
                
                # 快速查找并点击删除按钮
                delete_found = False
                delete_selectors = [
                    "//*[contains(text(), '删除此操作员')]",
                    "//button[contains(text(), '删除')]",
                    "//a[contains(text(), '删除')]",
                    "//*[contains(@onclick, 'delete')]",
                    "//*[contains(@class, 'delete')]"
                ]
                
                for selector in delete_selectors:
                    try:
                        delete_btn = self.driver.find_element(By.XPATH, selector)
                        if delete_btn.is_displayed():
                            delete_btn.click()
                            delete_found = True
                            time.sleep(0.3)  # 短暂等待弹窗
                            break
                    except:
                        continue
                
                if not delete_found:
                    return False
                
                # 快速确认删除
                confirm_selectors = [
                    "//*[contains(text(), '是，删除操作员')]",
                    "//button[contains(text(), '是')]",
                    "//button[contains(text(), '确认')]",
                    "//button[contains(text(), '确定')]",
                    "//*[contains(@class, 'confirm')]"
                ]
                
                for selector in confirm_selectors:
                    try:
                        confirm_btn = self.driver.find_element(By.XPATH, selector)
                        if confirm_btn.is_displayed():
                            confirm_btn.click()
                            time.sleep(0.5)  # 减少等待时间
                            return True
                    except:
                        continue
                
                # 如果没找到确认按钮，尝试按回车
                try:
                    from selenium.webdriver.common.action_chains import ActionChains
                    actions = ActionChains(self.driver)
                    actions.send_keys(Keys.ENTER).perform()
                    time.sleep(0.5)
                    return True
                except:
                    pass
                
                return False
            else:
                return False
                
        except Exception as e:
            return False
    
    def smart_wait(self, timeout=3):
        """
        智能等待页面加载完成
        """
        try:
            # 等待文档加载完成
            self.short_wait.until(
                lambda driver: self.driver.execute_script("return document.readyState") == "complete"
            )
            time.sleep(0.5)  # 额外短暂等待
        except:
            time.sleep(timeout)  # 如果智能等待失败，使用固定等待
    
    def execute_automation(self, start_url):
        try:
            # 打开起始页面
            self.driver.get(start_url)
            time.sleep(1)
            
            # 登录系统
            if not self.login():
                return
            
            total_records = len(self.data)
            start_time = time.time()
            
            for i, text in enumerate(self.data, 1):
                # 步骤1: 快速搜索
                search_result = self.find_and_paste_text(str(text))
                
                if search_result == "no_results":
                    continue
                elif not search_result:
                    continue
                
                # 步骤2: 快速点击超链接
                if not self.fast_click_by_text(str(text)):
                    continue
                
                # 步骤3: 快速检查并删除
                self.fast_check_and_delete()
            
            
        except Exception as e:
            print(f"自动化执行过程中出错: {e}")
        finally:
            message = (f"所有用户已删除！")
            log.insert("insert", message)
            self.driver.quit()
    
    def close(self):
        """关闭浏览器"""
        self.driver.quit()

def Delete_Voice_user():

    log.delete(1.0, tk.END)
    START_URL = "http://tcgiapp1wapp013:9090/VoiceConsole/core/search/result.action"  # 替换为你的目标网站URL
    DRIVER_PATH = None  # 如果Chrome驱动不在PATH中，请指定路径
    
    try:
        data = None
        data = sql.get("1.0", "end-1c")
        if data =='':
            message = (f"请输入条件！\n")
            log.insert("insert", message)
        else:
            automation = ExcelDeleteAutomation(DRIVER_PATH)
            automation.execute_automation(START_URL)
       
    except KeyboardInterrupt:
        print("\n用户中断程序")
    except Exception as e:
        print(f"程序执行出错: {e}")


class ExcelADDAutomation:
    def __init__(self, driver_path=None):
        # 读取Excel数据到缓存
        self.data_cache = self.read_excel_data()
        self.current_index = 0
        
        # 设置浏览器选项
        options = webdriver.ChromeOptions()
        options.add_experimental_option("prefs", {
            "credentials_enable_service": False,
            "profile.password_manager_enabled": False,
        })
        options.add_argument('--disable-blink-features=AutomationControlled')
        options.add_experimental_option("excludeSwitches", ["enable-automation"])
        options.add_experimental_option('useAutomationExtension', False)
        options.add_argument('--disable-extensions')
        options.add_argument('--no-sandbox')
        options.add_argument('--disable-dev-shm-usage')
        options.add_argument('--disable-gpu')
        
        # 初始化驱动
        try:
            if driver_path:
                self.driver = webdriver.Chrome(executable_path=driver_path, options=options)
            else:
                self.driver = webdriver.Chrome(options=options)
        except Exception as e:
            message = f"浏览器驱动初始化失败: {e}\n"
            log.insert("insert", message)
            raise
        
        # 隐藏自动化特征
        self.driver.execute_script("Object.defineProperty(navigator, 'webdriver', {get: () => undefined})")
        
        self.wait = WebDriverWait(self.driver, 10)
        self.short_wait = WebDriverWait(self.driver, 3)
    
    def read_excel_data(self):
        try:
            try:
                df = pd.read_excel('data.xlsx', header=None)
           
            except Exception as e:
                message = f"读取Excel文件失败: {e}\n"
                log.insert("insert", message)
                return []
            
            if df.empty:
                message = "Excel文件为空\n"
                log.insert("insert", message)
                return []
            
            # 确保数据有两列
            if len(df.columns) < 2:
                message = f"Excel文件需要至少2列数据，当前只有{len(df.columns)}列\n"
                log.insert("insert", message)
                return []
            
            # 转换为列表，确保处理所有行
            data_list = []
            for i, row in df.iterrows():
                if len(row) >= 2:
                    # 处理可能的NaN值
                    col1 = str(row.iloc[0]) if pd.notna(row.iloc[0]) else ""
                    col2 = str(row.iloc[1]) if pd.notna(row.iloc[1]) else ""
                    if col1 and col2:  # 只添加非空行
                        data_list.append([col1, col2])
            
            message = f"成功读取Excel数据，共 {len(data_list)} 条有效记录\n"
            log.insert("insert", message)
            return data_list
        except Exception as e:
            message = f"读取Excel数据失败: {e}\n"
            log.insert("insert", message)
            return []
    
    def has_next_record(self):
        """检查是否还有下一条记录"""
        return self.current_index < len(self.data_cache)
    
    def get_next_record(self):
        """获取下一条记录并移动指针"""
        if self.has_next_record():
            record = self.data_cache[self.current_index]
            self.current_index += 1
            return record
        return None
    
    def get_total_records(self):
        """获取总记录数"""
        return len(self.data_cache)
    
    def get_current_progress(self):
        """获取当前进度"""
        return self.current_index, len(self.data_cache)
    
    def login(self):
        """登录系统"""
        try:
            self.driver.get("http://tcgiapp1wapp013:9090/VoiceConsole/login.action")
            self.driver.maximize_window()
            
            # 使用显式等待
            wait = WebDriverWait(self.driver, 10)
            
            # 定位登录表单元素
            username = wait.until(EC.presence_of_element_located((By.NAME, "j_username")))
            password = wait.until(EC.presence_of_element_located((By.NAME, "j_password")))
        
            # 直接使用Selenium输入
            username.send_keys(VC["username"])
            password.send_keys(VC["password"])
            password.submit()  # 自动提交表单
            
            # 等待页面加载
            time.sleep(1)
            return True     
        except Exception as e:
            message = f"登录过程中出错: {e}\n"
            log.insert("insert", message)
            return False
    
    def navigate_to_create_operator(self):
        """导航到创建操作员页面"""
        try:
            # 直接跳转到操作员列表页面
            self.driver.get("http://tcgiapp1wapp013:9090/VoiceConsole/core/operator/list.action")
            time.sleep(1)
            
            # 点击"创建新操作员"按钮
            create_btn = self.wait.until(EC.element_to_be_clickable(
                (By.XPATH, "//a[contains(text(), '创建新操作员')]")
            ))
            create_btn.click()
            time.sleep(1)
            return True
        except Exception as e:
            message = f"导航到创建操作员页面失败: {e}\n"
            log.insert("insert", message)
            return False
    
    def select_task_package(self):
        try:
            task_select_element = self.wait.until(
                EC.presence_of_element_located((By.XPATH, "//select"))
            )
            
            select = Select(task_select_element)
            
            # 先尝试通过可见文本选择
            try:
                select.select_by_visible_text("V645_CN_20150624")
                return True
            except:
                pass            
        except Exception as e:
            message = f"  ✗ 选择任务包失败: {e}\n"
            log.insert("insert", message)
            return False
    
    def fill_operator_form(self, operator_name, operator_id):
        try:
            name_field = self.wait.until(EC.presence_of_element_located(
                (By.NAME, "operator.name")
            ))
            name_field.clear()
            name_field.send_keys(str(operator_name))
            
            # 2. 填写操作员ID
            id_field = self.wait.until(EC.presence_of_element_located(
                (By.NAME, "operator.operatorFuncId")
            ))
            id_field.clear()
            id_field.send_keys(str(operator_id))
            
            # 3. 填写朗读名称
            reading_name_field = self.wait.until(EC.presence_of_element_located(
                (By.NAME, "operator.spokenName")
            ))
            reading_name_field.clear()
            reading_name_field.send_keys(str(operator_id))
            # 4. 勾选自动生成操作员号
            try:
                auto_checkbox = self.driver.find_element(By.ID, "generateId")
                if not auto_checkbox.is_selected():
                    auto_checkbox.click()
            except:
                message = "  ⓘ 勾选复选框操作失败\n"
                log.insert("insert", message)
            
            # 5. 选择任务包
            if not self.select_task_package():
                message = "  ✗ 选择任务包失败\n"
                log.insert("insert", message)
                return False
            
            # 6. 选择操作员组（最小的非PS组）
            if not self.select_operator_group():
                message = "  ⓘ 选择操作员组时出现问题\n"
                log.insert("insert", message)
            
            return True
            
        except Exception as e:
            message = f"填写操作员表单失败: {e}\n"
            log.insert("insert", message)
            return False
    
    def select_operator_group(self):
        """选择操作员组"""
        try:
            # 查找所有操作员组的复选框
            checkboxes = self.driver.find_elements(By.XPATH, "//input[@name='addTeamIds']")
            group_rows = self.driver.find_elements(By.XPATH, "//tr[contains(@class, 'even') or contains(@class, 'odd')]")
            
            smallest_group_checkbox = None
            smallest_size = float('inf')
            selected_group_name = ""
            
            for i, checkbox in enumerate(checkboxes):
                try:
                    # 获取对应的行
                    if i < len(group_rows):
                        row = group_rows[i]
                        cells = row.find_elements(By.TAG_NAME, "td")
                        
                        if len(cells) >= 5:
                            group_name = cells[1].text.strip()  # 第二列是组名
                            group_size_text = cells[4].text.strip()  # 第五列是组大小
                            
                            # 跳过包含PS的组
                            if "PS" in group_name.upper():
                                continue
                            
                            # 提取数字
                            import re
                            size_match = re.search(r'\d+', group_size_text)
                            group_size = int(size_match.group()) if size_match else 0
                            
                            if group_size < smallest_size:
                                smallest_group_checkbox = checkbox
                                smallest_size = group_size
                                selected_group_name = group_name
                except:
                    continue
            
            if smallest_group_checkbox:
                if not smallest_group_checkbox.is_selected():
                    smallest_group_checkbox.click()
                return True
            else:
                # 如果没有找到非PS组，选择第一个可用的组
                if checkboxes:
                    first_checkbox = checkboxes[0]
                    if not first_checkbox.is_selected():
                        first_checkbox.click()
                    message = "  ✓ 已选择第一个可用组（未找到非PS组）\n"
                    log.insert("insert", message)
                    return True
                else:
                    message = "  ✗ 未找到任何可用的操作员组\n"
                    log.insert("insert", message)
                    return False
                    
        except Exception as e:
            message = f"  ✗ 选择操作员组失败: {e}\n"
            log.insert("insert", message)
            return False
    
    def save_and_create_another(self):
        """保存并创建另一个操作员"""
        try:
            # 查找保存按钮
            save_buttons = [
                "//a[@id='saveCreateIndex']",  # 保存并创建另一个按钮
                "//a[contains(text(), '保存并创建另一个')]",
                "//input[@id='form1.submit1']",  # 创建操作员按钮
                "//a[contains(text(), '创建操作员')]"
            ]
            
            save_btn = None
            for xpath in save_buttons:
                try:
                    save_btn = self.wait.until(EC.element_to_be_clickable((By.XPATH, xpath)))
                    break
                except:
                    continue
            
            if not save_btn:
                message = "  ✗ 未找到保存按钮\n"
                log.insert("insert", message)
                return False
            
            # 优先使用"保存并创建另一个"按钮
            if "saveCreateIndex" in save_btn.get_attribute("id") or "保存并创建另一个" in save_btn.text:
                save_btn.click()
            else:
                # 如果没有找到"保存并创建另一个"，使用普通创建按钮
                save_btn.click()
            
            time.sleep(1)  # 等待操作完成
            
            # 检查是否出现重复操作员错误
            if self.check_duplicate_error():
                return "duplicate"
            
            # 检查是否成功保存并返回到创建页面
            if self.check_successful_save():
                return True
            else:
                return False
                
        except Exception as e:
            message = f"保存操作失败: {e}\n"
            log.insert("insert", message)
            return False
    
    def check_duplicate_error(self):
        """检查重复操作员错误"""
        try:
            page_text = self.driver.find_element(By.TAG_NAME, "body").text
            duplicate_indicators = [
                "操作员姓名 已经存在",
                "操作员姓名 MENG 已经存在",
                "already exists", 
                "重复", 
                "duplicate"
            ]
            
            for indicator in duplicate_indicators:
                if indicator in page_text:
                    message = f"  ⓘ 用户已存在: {indicator}\n"
                    log.insert("insert", message)
                    return True
            return False
        except:
            return False
    
    def check_successful_save(self):
        """检查是否成功保存"""
        try:
            # 检查是否返回到创建页面（查看是否有操作员姓名字段）
            name_field = self.driver.find_element(By.NAME, "operator.name")
            is_empty = name_field.get_attribute("value") == ""
            if is_empty:
                message = "  ✓ 成功保存并返回创建页面\n"
                log.insert("insert", message)
            return is_empty
        except:
            # 如果找不到姓名字段，可能保存成功但页面跳转
            try:
                # 检查当前URL是否还在创建页面
                if "create" in self.driver.current_url or "list" in self.driver.current_url:
                    message = "  ✓ 保存成功，页面已跳转\n"
                    log.insert("insert", message)
                    return True
            except:
                pass
            return False
    
    def execute_automation(self):
        """执行自动化流程"""
        try:
            if not self.data_cache:
                message = "Excel中没有数据或读取失败\n"
                log.insert("insert", message)
                return
            # 登录系统
            if not self.login():
                message = "登录失败，终止程序\n"
                log.insert("insert", message)
                return
            
            # 导航到创建操作员页面
            if not self.navigate_to_create_operator():
                message = "导航到创建页面失败，终止程序\n"
                log.insert("insert", message)
                return
            
            success_count = 0
            duplicate_count = 0
            error_count = 0
            
            # 逐条处理记录
            while self.has_next_record():
                current_index, total = self.get_current_progress()
                record = self.get_next_record()
                
                if record and len(record) >= 2:
                    operator_name = str(record[0])  # 第一列：操作员姓名
                    operator_id = str(record[1])    # 第二列：操作员ID和朗读名称
                    
                    # 填写表单
                    if self.fill_operator_form(operator_name, operator_id):
                        # 保存操作员
                        result = self.save_and_create_another()
                        
                        if result == "duplicate":
                            duplicate_count += 1
                        elif result:
                            success_count += 1
                        else:
                            error_count += 1
                    else:
                        error_count += 1
                else:
                    error_count += 1
            message = f"\n处理完成:\n"
            message += f"成功添加: {success_count} 个操作员\n"
            message += f"跳过重复: {duplicate_count} 个操作员\n"
            message += f"处理失败: {error_count} 个操作员\n"
            log.insert("insert", message)
            
        except Exception as e:
            message = f"自动化执行过程中出错: {e}\n"
            log.insert("insert", message)
        finally:
            try:
                self.driver.quit()
            except:
                pass

def Add_Voice_user():
    """添加语音用户主函数"""
    log.delete(1.0, tk.END)
    
    try:
        automation = ExcelADDAutomation()
        automation.execute_automation()
        
    except KeyboardInterrupt:
        message = "用户中断程序\n"
        log.insert("insert", message)
    except Exception as e:
        message = f"程序执行出错: {e}\n"
        log.insert("insert", message)

def Export():

    try:

        cnxn = pyodbc.connect(con_string)

        cur = cnxn.cursor()

        target = path.get()
        filename = file_name.get()

        filetype = '.xlsx'

        file = filename + filetype

        if cur:

            message = f'{dt.now()}: DB连接成功 !\n'
            log.insert("insert", message)

        else:

            message = f'{dt.now()}: DB连接失败 !\n'
            log.insert("insert", message)


        SQL = sql.get("1.0", "end-1c")


        data = cur.execute(SQL).fetchall()

        fields = [field[0] for field in cur.description]

        f_name = os.path.join(target, file)

        if len(data):

            rows = []

            for d in data:
                rows.append(list(d))


            df = pd.DataFrame(data=rows, columns=fields)
            for field in fields: 
                df[field] = df[field].apply(bytes_to_string)
            df = process_data_with_sql(df)
            df.to_excel(f_name, index=False)
            cur.close()


            message = f'{dt.now()}: {f_name} 已生成 !\n'
            log.insert("insert", message)

            # messagebox.showinfo("成功", "请查看数据")

        else:

            cur.close()

            message = f'{dt.now()}: 无数据 !\n'
            log.insert("insert", message)

            # messagebox.showinfo("成功", "无数据")


    except Exception as e:

        message = f'{dt.now()}: {e}\n'
        log.insert("insert", message)



class DatabaseViewer(tk.Frame):

    def __init__(self, master, fields, rows):
        super().__init__(master)
        self.master = master
        self.fields = fields
        self.rows = rows
        self.canvas = tk.Canvas(self, width=850, height=450)

        self.vscrollbar = ttk.Scrollbar(self, orient=tk.VERTICAL, command=self.canvas.yview)
        self.hscrollbar = ttk.Scrollbar(self, orient=tk.HORIZONTAL, command=self.canvas.xview)
        self.canvas.configure(yscrollcommand=self.vscrollbar.set, xscrollcommand=self.hscrollbar.set)

        self.frame = tk.Frame(self.canvas)
        self.canvas.create_window((0, 0), window=self.frame, anchor="nw", tags="frame")

        self.vscrollbar.grid(row=0, column=1, sticky="ns")
        self.hscrollbar.grid(row=1, column=0, sticky="we")
        self.canvas.grid(row=0, column=0, sticky="wens")

        self.canvas.bind("<MouseWheel>", self.on_mousewheel)

        # 设置字体和字号
        font_style = ("Calibri", 12)

        # 表头部分
        for j, value in enumerate(fields):
            label = tk.Label(self.frame, text=str(value).strip(), borderwidth=1, relief="solid", bg='#DDEBF7', fg='black', font=('Arial', 10, 'bold'))
            label.grid(row=0, column=j, sticky="nsew")

        # 内容部分
        for i, row in enumerate(rows):
            for j, value in enumerate(row):
                label = tk.Label(self.frame, text=str(value).strip(), borderwidth=1, relief="solid", bg='white', fg='black', font=font_style)
                label.grid(row=i+1, column=j, sticky="nsew")
            if i > 35:
                break

        # 调整列宽和行高
        for col in range(len(fields)):
            self.frame.columnconfigure(col, weight=1)
        for row in range(len(rows) + 1):
            self.frame.rowconfigure(row, weight=1)

        self.frame.update_idletasks()
        self.canvas.configure(scrollregion=self.canvas.bbox("all"))

    def on_mousewheel(self, event):
        self.canvas.yview_scroll(int(-1 * (event.delta / 120)), "units")
        
def Lable_His():
    label_log.delete(1.0, tk.END)
    try:
        # 获取特定条件输入框的值
        # specific_condition_raw = specific_input.get("1.0", "end-1c").strip()
        # specific_condition = [line.strip() for line in specific_condition_raw.splitlines()]
        # specific_condition = [line.replace(" ", "") for line in specific_condition]



        # 数据库连接
        jls_extract_var = con_string
        cnxn = pyodbc.connect(jls_extract_var)
        cur = cnxn.cursor()

        if cur:
            message = f'{dt.now()}: DB连接成功 !\n'
            label_log.insert("insert", message)
        else:
            message = f'{dt.now()}: DB连接失败 !\n'
            label_log.insert("insert", message)
        # # 特定条件查询
        # if specific_condition:
        #     placeholders = ', '.join(['?'] * len(specific_condition))
        #     # 使用特定格式的SQL语句
        #     SQL =f'''
        #     SELECT C.GRGRPCDE, C.LOSYSLOC, A.PRSTYCDE, A.PRCOLCDE, A.INEXTSIZ
        #     FROM CTXPRWDTA.INRDTA A,CTXPRWDTA.INRHLA B,CTXPRWDTA.GRRALO C
        #     WHERE A.INHLDCDE = B.INHLDCDE AND
        #     trim(A.PRSTYCDE)||'-'||trim(A.PRCOLCDE)||'-'||trim(A.INEXTSIZ) in ({placeholders})
        #     AND A.INHLDCDE LIKE '00000000003%'
        #     AND B.LOLOCSGT= C.LOLOCSGT
        #     AND (C.LOSYSLOC LIKE '35%' OR C.LOSYSLOC LIKE '38%' OR C.LOSYSLOC LIKE '39%')
        #     GROUP BY C.GRGRPCDE, C.LOSYSLOC, A.PRSTYCDE, A.PRCOLCDE, A.INEXTSIZ
        #     ORDER BY A.PRSTYCDE, A.PRCOLCDE, A.INEXTSIZ
        #     '''

        #     # 执行SQL语句，将specific_condition作为参数传递
        #     data = cur.execute(SQL, tuple(specific_condition)).fetchall()
        #     if data:
        #         fields = [field[0] for field in cur.description]
        #         display_results(data, fields)
        #     else:
        #         # 特定条件查询无结果，尝试箱号查询
        #         #pass
        #         #message = f'{dt.now()}: {c.strip()} 没有查到历史库位 !\n'
        #         message = f'{dt.now()}: 该SKU没有查到历史库位 !\n'
        #         label_log.insert("insert", message)
                
        # 箱号查询
        crtns = sorted(set(c.strip() for c in carton.get("1.0", tk.END).split('\n') if c.strip()), 
               key=lambda x: list(dict.fromkeys(c.strip() for c in carton.get("1.0", tk.END).split('\n') if c.strip())).index(x))
        data = []
        for c in crtns:

            if query_type2.get() == "查询打印记录":
                SQL = SQL_DATA['label_history']['print_record']
            elif query_type2.get() == "CHECK OPEN WORK":
                 SQL = SQL_DATA['label_history']['check_open_work']
            elif query_type2.get() == "CHECK AUDIT WORK":
                SQL = SQL_DATA['label_history']['check_audit_work']
            elif query_type2.get() == "检查SKU是否是AMINUS":
                SQL = SQL_DATA['label_history']['check_sku_aminus']
            elif query_type2.get() == "QA解锁?":
                SQL = SQL_DATA['label_history']['qa_unlock']
            elif query_type2.get() == "SO查询":
                SQL = SQL_DATA['label_history']['so_query']
            elif query_type2.get() == "NFC":
                SQL = SQL_DATA['label_history']['nfc']
            elif query_type2.get() == "查询包装信息":
                SQL = SQL_DATA['label_history']['pack_info']
            elif query_type2.get() == "RSO":
                SQL = SQL_DATA['label_history']['rso']
            elif query_type2.get() == "查询历史库位":
                SQL = SQL_DATA['label_history']['history_location']
            elif query_type2.get() == "DP_area_info":
                    SQL = SQL_DATA['label_history']['dp_area_info']
            elif query_type2.get() == "Export inventory":
                params = {'CARTON_PLAN': c.strip()}
                SQLs = [
                    SQL_DATA['label_history']['export_inventory_q1'].format(**params),
                    SQL_DATA['label_history']['export_inventory_q2'].format(**params),
                    SQL_DATA['label_history']['export_inventory_q3'].format(**params),
                    SQL_DATA['label_history']['export_inventory_q4'].format(**params)
                ]
                # Create a new Excel writer object using xlsxwriter
                excel_filename = f"inventory.xlsx"
                workbook = xlsxwriter.Workbook(excel_filename)
                worksheet = workbook.add_worksheet('Inventory')

                # Define a format for centering text
                center_format = workbook.add_format({'align': 'center', 'valign': 'vcenter'})

                start_row = 0
                for i, sql in enumerate(SQLs):
                    formatted_sql = sql.format(**params)
                    rows = cur.execute(formatted_sql).fetchall()
                    fields = [field[0] for field in cur.description]

                    # Write headers
                    worksheet.write_row(start_row, 0, fields, center_format)
                    start_row += 1

                    # Write the data rows
                    for row in rows:
                        worksheet.write_row(start_row, 0, row, center_format)
                        start_row += 1

                    # Leave two empty rows between queries
                    start_row += 2

                # Close the workbook
                workbook.close()

                message = f'{dt.now()}: 数据已导出至 {excel_filename}!\n'
                label_log.insert("insert", message)
                return    
            else:
                SQL = SQL_DATA['label_history']['default_location']
            SQL = SQL.replace('CARTON_PLAN', f"{c.strip()}")
            rows = cur.execute(SQL).fetchall()
            if len(rows):
                for r in rows:
                    data.append(r)
                fields = [field[0] for field in cur.description]
                message = f'{dt.now()}: {c.strip()} 有数据 !\n'
                label_log.insert("insert", message)
            else:
                message = f'{dt.now()}: {c.strip()} 无数据 !\n'
                label_log.insert("insert", message)

        # 关闭游标
        cur.close()

        # 如果有数据，显示结果
        if data:
            can = tk.Tk()
            display_results(data, fields,can)
            export_button2 = ttk.Button(can, text="导出")
            export_button2.grid(row=0, column=1, columnspan=1, sticky='n')
            export_button2.config(command=lambda: export_data_overview(data, fields, can))

    except Exception as e:
        message = f'{dt.now()}: {e}\n'
        label_log.insert("insert", message)

def display_results(data, fields, can):
    for widget in can.winfo_children():
        widget.destroy()

    # 数据量检查，超过1000行启用分页显示
    if len(data) > 1000:
        create_paginated_view(can, data, fields)
    else:
        # 小数据量直接显示
        converted_data = []
        for row in data:
            converted_row = [bytes_to_string(item) for item in row]
            converted_data.append(converted_row)
        
        database_viewer = DatabaseViewer(master=can, fields=fields, rows=converted_data)
        database_viewer.grid(row=0, column=0, sticky="nsew")
    
    # 设置窗口尺寸和位置
    screen_width = can.winfo_screenwidth()
    screen_height = can.winfo_screenheight()
    window_width = 1000
    window_height = 800
    pos_right = int(screen_width / 2 - window_width / 2)
    pos_down = int(screen_height / 2 - window_height / 2)
    can.geometry(f"{window_width}x{window_height}+{pos_right}+{pos_down}")

def create_paginated_view(can, data, fields, page_size=500):
    """创建分页显示界面"""
    # 创建主框架
    main_frame = ttk.Frame(can)
    main_frame.grid(row=0, column=0, sticky="nsew")
    can.grid_rowconfigure(0, weight=1)
    can.grid_columnconfigure(0, weight=1)
    
    # 分页管理器
    class PaginationManager:
        def __init__(self, data, page_size):
            self.data = data
            self.page_size = page_size
            self.current_page = 0
            self.total_pages = (len(data) + page_size - 1) // page_size
            
        def get_page_data(self, page_num):
            start_idx = page_num * self.page_size
            end_idx = min(start_idx + self.page_size, len(self.data))
            return self.data[start_idx:end_idx]
    
    paginator = PaginationManager(data, page_size)
    
    # 顶部信息栏
    top_frame = ttk.Frame(main_frame)
    top_frame.grid(row=0, column=0, columnspan=3, sticky="ew", pady=10)
    
    info_label = ttk.Label(top_frame, 
                          text=f"数据量较大 ({len(data):,} 行，共{paginator.total_pages}页)，已启用分页显示",
                          font=("Arial", 10, "bold"))
    info_label.pack(side="left", padx=10)
    
    export_btn = ttk.Button(top_frame, text="📊 导出全部数据到Excel", 
                          command=lambda: export_large_data(data, fields, can))
    export_btn.pack(side="right", padx=10)
    
    # 分页控制栏
    control_frame = ttk.Frame(main_frame)
    control_frame.grid(row=1, column=0, columnspan=3, sticky="ew", pady=5)
    
    # 分页按钮和标签
    ttk.Button(control_frame, text="◀◀ 首页", 
              command=lambda: change_page(0)).pack(side="left", padx=5)
    
    ttk.Button(control_frame, text="◀ 上一页", 
              command=lambda: change_page(paginator.current_page - 1)).pack(side="left", padx=5)
    
    page_label = ttk.Label(control_frame, text="", font=("Arial", 10))
    page_label.pack(side="left", padx=20)
    
    ttk.Button(control_frame, text="下一页 ▶", 
              command=lambda: change_page(paginator.current_page + 1)).pack(side="right", padx=5)
    
    ttk.Button(control_frame, text="末页 ▶▶", 
              command=lambda: change_page(paginator.total_pages - 1)).pack(side="right", padx=5)
    
    # 跳转页面输入框
    jump_frame = ttk.Frame(control_frame)
    jump_frame.pack(side="right", padx=20)
    
    ttk.Label(jump_frame, text="跳转到:").pack(side="left")
    page_entry = ttk.Entry(jump_frame, width=5)
    page_entry.pack(side="left", padx=5)
    ttk.Button(jump_frame, text="跳转", 
              command=lambda: jump_to_page()).pack(side="left")
    
    # 数据显示区域
    display_frame = ttk.Frame(main_frame)
    display_frame.grid(row=2, column=0, columnspan=3, sticky="nsew")
    
    main_frame.grid_rowconfigure(2, weight=1)
    main_frame.grid_columnconfigure(0, weight=1)
    
    def update_display():
        """更新数据显示"""
        # 清除现有显示
        for widget in display_frame.winfo_children():
            widget.destroy()
        
        # 获取当前页数据
        page_data = paginator.get_page_data(paginator.current_page)
        start_idx = paginator.current_page * paginator.page_size + 1
        end_idx = min(start_idx + len(page_data) - 1, len(data))
        
        # 更新页面标签
        page_label.config(text=f"第 {paginator.current_page + 1}/{paginator.total_pages} 页 | "
                              f"显示行: {start_idx}-{end_idx} | 总行数: {len(data):,}")
        
        # 转换数据格式
        converted_data = []
        for row in page_data:
            converted_row = [bytes_to_string(item) for item in row]
            converted_data.append(converted_row)
        
        # 创建表格显示
        database_viewer = DatabaseViewer(master=display_frame, fields=fields, rows=converted_data)
        database_viewer.grid(row=0, column=0, sticky="nsew")
        
        display_frame.grid_rowconfigure(0, weight=1)
        display_frame.grid_columnconfigure(0, weight=1)
    
    def change_page(page_num):
        """切换页面"""
        if 0 <= page_num < paginator.total_pages:
            paginator.current_page = page_num
            update_display()
    
    def jump_to_page():
        """跳转到指定页面"""
        try:
            page_num = int(page_entry.get()) - 1
            if 0 <= page_num < paginator.total_pages:
                change_page(page_num)
            else:
                tk.messagebox.showwarning("警告", f"页码必须在 1 到 {paginator.total_pages} 之间")
        except ValueError:
            tk.messagebox.showerror("错误", "请输入有效的页码数字")
    
    # 初始显示第一页
    update_display()

def export_large_data(data, fields, can):
    """导出大量数据到Excel"""
    try:
        # 显示导出进度
        progress_window = tk.Toplevel(can)
        progress_window.title("导出数据")
        progress_window.geometry("300x100")
        progress_window.transient(can)
        progress_window.grab_set()
        
        progress_label = ttk.Label(progress_window, text="正在导出数据，请稍候...")
        progress_label.pack(pady=10)
        
        progress_var = tk.DoubleVar()
        progress_bar = ttk.Progressbar(progress_window, variable=progress_var, maximum=100)
        progress_bar.pack(fill="x", padx=20, pady=5)
        
        progress_window.update()
        
        # 分批处理数据以避免内存问题
        chunk_size = 1000
        converted_data = []
        
        for i in range(0, len(data), chunk_size):
            chunk = data[i:i + chunk_size]
            for row in chunk:
                converted_row = [bytes_to_string(item) for item in row]
                converted_data.append(converted_row)
            
            # 更新进度
            progress = min((i + len(chunk)) / len(data) * 100, 100)
            progress_var.set(progress)
            progress_window.update()
        
        # 创建DataFrame
        df = pd.DataFrame([list(row) for row in converted_data], columns=fields)
        df = process_data_with_sql(df)
        
        # 数据清洗
        for col in df.select_dtypes(include='object').columns:
            if df[col].dtype == 'O':
                df[col] = df[col].astype(str).str.strip().str.replace(r'\s+', '', regex=True)
                df[col] = df[col].astype(str).str.replace('None', '', regex=True)
        
        # 获取桌面路径
        desktop_path = os.path.join(os.path.expanduser("~"), 'Desktop')
        file_name = f'data.xlsx'
        file_path = os.path.join(desktop_path, file_name)
        
        # 写入Excel
        df.to_excel(file_path, index=False)
        
        progress_window.destroy()
        tk.messagebox.showinfo("导出成功", 
                              f'文件已保存至桌面: {file_name}\n'
                              f'数据量: {len(data):,} 行 × {len(fields)} 列')
        can.destroy()                     
        
    except Exception as e:
        if 'progress_window' in locals():
            progress_window.destroy()
        tk.messagebox.showerror("导出失败", f"导出过程中出现错误:\n{str(e)}")

def export_data_overview(data, fields, can):
    """导出数据概览（兼容大小数据量）"""
    if len(data) > 10000:
        # 大数据量使用专门的导出函数
        export_large_data(data, fields, can)
    else:
        # 小数据量使用原有逻辑
        converted_data = []
        for row in data:
            converted_row = [bytes_to_string(item) for item in row]
            converted_data.append(converted_row)
        
        df2 = pd.DataFrame([list(row) for row in converted_data], columns=fields)
        
        # 去除所有字符串列的前后空格和内部空格
        for col in df2.select_dtypes(include='object').columns:
            if df2[col].dtype == 'O':
                df2[col] = df2[col].astype(str).str.strip().str.replace(r'\s+', '', regex=True)
                df2[col] = df2[col].astype(str).str.replace('None', '', regex=True)
        
        df2 = process_data_with_sql(df2)

        # 获取桌面路径
        desktop_path = os.path.join(os.path.expanduser("~"), 'Desktop')
        file_name = f'data.xlsx'
        file_path = os.path.join(desktop_path, file_name)

        # 写入Excel文件
        df2.to_excel(file_path, index=False)
        tk.messagebox.showinfo("导出成功", f'文件已保存至桌面: {file_name}')
        can.destroy()

########################################################
#  frame area
########################################################

option_frame = tk.Frame(master=data_export,bd=2)


# oPen = tk.Button(
#     master=option_frame,

#     text="打开",

#     command=Open_File
#     )


# copy = tk.Button(
#     master=option_frame,

#     text="复制",

#     command=Copy
#     )


# paste = tk.Button(
#     master=option_frame,

#     text="粘贴",

#     command=Paste
#     )


clear = tk.Button(
    master=option_frame,

    text="清屏",

    command=Clear
    )


formate = tk.Button(
    master=option_frame,

    text="格式化SQL",

    command=Formate
    )

GI_status = tk.Button(
    master=option_frame,

    text="GI_status_check",

    command=GI_status
)

health_check = tk.Button(
    master=option_frame,

    text="health_check",

    command=health_check
)

Delete_Voice_user = tk.Button(
    master=option_frame,

    text="Delete_Voice_user",

    command=Delete_Voice_user
)

Add_Voice_user = tk.Button(
    master=option_frame,

    text="Add_Voice_user",

    command=Add_Voice_user
)
# oPen.grid(row=0, column=0,padx=2,pady=2)

# copy.grid(row=0, column=1,padx=2,pady=2)

# paste.grid(row=0, column=2,padx=2,pady=2)

clear.grid(row=0, column=3,padx=2,pady=2)

formate.grid(row=0, column=4,padx=2,pady=2)

GI_status.grid(row=0, column=5,padx=2,pady=2)

health_check.grid(row=0, column=6,padx=2,pady=2)

Delete_Voice_user.grid(row=0, column=7,padx=2,pady=2)

Add_Voice_user.grid(row=0, column=8,padx=2,pady=2)
########

sql_frame = tk.Frame(master=data_export,bd=2)


sql = tk.Text(

    master=sql_frame
    ,fg="black"

    ,bg="white"

    ,width=100
    ,height=10
    )

# text.pack(fill=tk.BOTH,expand=tk.TRUE)


sql.pack(side=tk.LEFT,padx=2,pady=2,fill=tk.BOTH, expand=tk.TRUE)


scrollbar = ttk.Scrollbar(sql_frame, command=sql.yview)

scrollbar.pack(side=tk.RIGHT, fill=tk.Y)


sql.configure(yscrollcommand=scrollbar.set)

########

file_frame = tk.Frame(master=data_export,bd=2)


filename_txt = tk.Label(
    master=file_frame

    # text=var_lbl,

    ,text ='请输入文件名: '

    # ,fg="white"
    # ,bg="black"

    ,width=10

    ,height=2
    )


file_name = tk.Entry(
    master=file_frame
    ,fg="black"

    ,bg="white"

    ,width=15
    )


# filetype_txt = tk.Label(
#     master=file_frame

#     # text=var_lbl,

#     ,text ='请选择格式:'

#     # ,fg="white"
#     # ,bg="black"

#     ,width=12

#     ,height=2
#     )


# file_formats = ["xlsx", "xls"]

# file_type = ttk.Combobox(file_frame, values=file_formats, state="readonly", width= 5, height= 1)



filename_txt.pack(side=tk.LEFT,padx=2,pady=2)

file_name.pack(side=tk.LEFT,padx=2,pady=2)

# filetype_txt.pack(side=tk.LEFT,padx=2,pady=2)

# file_type.pack(side=tk.LEFT,padx=2,pady=2)

########

folder_txt_frame = tk.Frame(master=data_export,bd=2)


folder_txt = tk.Label(

    master=folder_txt_frame

    # text=var_lbl,

    ,text ='请选择目标文件夹: '

    # ,fg="white"
    # ,bg="black"

    ,width=14

    ,height=2
    )

folder_txt.pack(side=tk.LEFT)


########

path_frame = tk.Frame(master=data_export,bd=2)


path = tk.Entry(

    master=path_frame
    ,fg="black"

    ,bg="white"

    ,width=100
    )


export = tk.Button(

    master=path_frame,

    text="导出",

    command=Export
    )


browser = tk.Button(

    master=path_frame,

    text="浏览",

    command=Find_Path
    )


path.pack(side=tk.TOP,padx=2,pady=2)

export.pack(side=tk.RIGHT,padx=2,pady=2)

browser.pack(side=tk.RIGHT,padx=2,pady=2)

########


# 创建Shipment功能页面
shipment = ttk.Frame(notebook)

# 创建数字条件输入框
num_condition_label = ttk.Label(shipment, text="请输入计划号或shipment或packlist或run或箱号:")
num_condition_text = tk.Text(shipment, height=10, width=40)

# 创建字母条件输入框
alpha_condition_label = ttk.Label(shipment, text="请输入物料或SKU或ACCESS NUMBER:")
alpha_condition_text = tk.Text(shipment, height=10, width=40)

# 创建日志输出框
log_label_shipment = ttk.Label(shipment, text="Log:")
log_text_shipment = tk.Text(shipment, height=20, width=90)

# 创建滚动条
log_scrollbar_shipment = ttk.Scrollbar(shipment, command=log_text_shipment.yview)
log_text_shipment['yscrollcommand'] = log_scrollbar_shipment.set

# 创建导出按钮
export_button = ttk.Button(shipment, text="导出",command=lambda: export_data())


# 定位组件
num_condition_label.grid(row=0, column=0, sticky='w', pady=5, padx=5)
num_condition_text.grid(row=1, column=0, sticky='w', pady=5, padx=5)
alpha_condition_label.grid(row=0, column=1, sticky='w', pady=5, padx=5)
alpha_condition_text.grid(row=1, column=1, sticky='w', pady=5, padx=5)
log_label_shipment.grid(row=2, column=0, sticky='w', pady=5, padx=5)
log_text_shipment.grid(row=3, column=0, columnspan=2, sticky='nw', pady=5, padx=5)
log_scrollbar_shipment.grid(row=3, column=2, sticky="wns")
export_button.grid(row=4, column=0, columnspan=2, sticky='w', pady=5, padx=5)

num_condition_text.bind("<FocusIn>", lambda event: clear_other_inputs(event, [alpha_condition_text]))
alpha_condition_text.bind("<FocusIn>", lambda event: clear_other_inputs(event, [num_condition_text]))

def log_message(log_widget, message):
    log_widget.delete(1.0, tk.END)
    log_widget.insert(tk.END, f'{dt.now()}: {message}\n')

# 导出数据的函数

def execute_sql_and_check_results(cursor, sql, params):
    """执行SQL查询并检查结果是否为空"""
    cursor.execute(sql, params)
    results = cursor.fetchall()
    if not results:
        return None
    return results
def export_data():
    
    try:
        num_conditions = [c.strip() for c in num_condition_text.get("1.0", tk.END).split('\n') if c.strip()]
        alpha_conditions = [c.strip() for c in alpha_condition_text.get("1.0", tk.END).split('\n') if c.strip()]
        # 初始化SQL变量
        SQL = None
                # 连接数据库
        cnxn = pyodbc.connect(con_string)
        cur = cnxn.cursor()


        if cur:
            log_message(log_text_shipment,'DB连接成功 !')
        else:
            log_message(log_text_shipment,'DB连接失败 !')
        # 根据输入的条件类型和位数选择SQL语句
        if num_conditions:
            if len(num_conditions[0]) <= 4:
                placeholders = ', '.join(['?'] * len(num_conditions))
                SQL = SQL_DATA['export_data']['ap_shipping']
                file_prefix = "AP_Shipping"
                data = cur.execute(SQL, num_conditions).fetchall()
                # 获取字段名
                fields = [field[0] for field in cur.description]

                # 检查数据是否存在
                if not data:
                    log_message(log_text_shipment,'查询结果为空，没有数据可导出。')
                # 检查数据和字段名数量是否匹配
                elif len(data[0]) != len(fields):
                    log_message(log_text_shipment,'数据与字段名数量不匹配，跳过。')
                else:
                    # 创建DataFrame
                    df = pd.DataFrame([list(row) for row in data], columns=fields)
                    # 去除所有字符串列的前后空格和内部空格
                    for col in df.select_dtypes(include='object').columns:
                        if df[col].dtype == 'O':  # 'O' 表示 object 类型，即字符串或混合类型
                            df[col] = df[col].astype(str).str.strip().str.replace(r'\s+', '', regex=True)
                            df[col] = df[col].astype(str).str.replace('None', '', regex=True)
                    df = process_data_with_sql(df)

                    # 获取桌面路径
                    desktop_path = os.path.join(os.path.expanduser("~"), 'Desktop')
                    file_name = f'{file_prefix}.xlsx'
                    file_path = os.path.join(desktop_path, file_name)
                    df.to_excel(file_path, index=False)
                    log_message(log_text_shipment,'文件已保存至桌面 !')
                
                SQL = SQL_DATA['export_data']['fw_shipping']
                file_prefix = "FW_Shipping"
                data = cur.execute(SQL, num_conditions).fetchall()  # 修正：使用SQL1
                # 获取字段名
                fields = [field[0] for field in cur.description]

                # 检查数据是否存在
                if not data:
                    log_message(log_text_shipment,'查询结果为空，没有数据可导出。')
                # 检查数据和字段名数量是否匹配
                elif len(data[0]) != len(fields):
                    log_message(log_text_shipment,'数据与字段名数量不匹配，跳过。')
                else:
                    # 创建DataFrame
                    df = pd.DataFrame([list(row) for row in data], columns=fields)
                    # 去除所有字符串列的前后空格和内部空格
                    for col in df.select_dtypes(include='object').columns:
                        if df[col].dtype == 'O':  # 'O' 表示 object 类型，即字符串或混合类型
                            df[col] = df[col].astype(str).str.strip().str.replace(r'\s+', '', regex=True)
                            df[col] = df[col].astype(str).str.replace('None', '', regex=True)
                    df = process_data_with_sql(df)

                    # 获取桌面路径
                    desktop_path = os.path.join(os.path.expanduser("~"), 'Desktop')
                    file_name = f'{file_prefix}.xlsx'
                    file_path = os.path.join(desktop_path, file_name)
                    df.to_excel(file_path, index=False)
                    log_message(log_text_shipment,'文件已保存至桌面 !')
                
                SQL = SQL_DATA['export_data']['replenishment']
                file_prefix = "Replenishment"
                data = cur.execute(SQL, num_conditions).fetchall()
                # 获取字段名
                fields = [field[0] for field in cur.description]

                # 检查数据是否存在
                if not data:
                    log_message(log_text_shipment,'查询结果为空，没有数据可导出。')
                # 检查数据和字段名数量是否匹配
                elif len(data[0]) != len(fields):
                    log_message(log_text_shipment,'数据与字段名数量不匹配，跳过。')
                else:
                    # 创建DataFrame
                    df = pd.DataFrame([list(row) for row in data], columns=fields)
                    # 去除所有字符串列的前后空格和内部空格
                    for col in df.select_dtypes(include='object').columns:
                        if df[col].dtype == 'O':  # 'O' 表示 object 类型，即字符串或混合类型
                            df[col] = df[col].astype(str).str.strip().str.replace(r'\s+', '', regex=True)
                            df[col] = df[col].astype(str).str.replace('None', '', regex=True)
                    df = process_data_with_sql(df)

                    # 获取桌面路径
                    desktop_path = os.path.join(os.path.expanduser("~"), 'Desktop')
                    file_name = f'{file_prefix}.xlsx'
                    file_path = os.path.join(desktop_path, file_name)
                    df.to_excel(file_path, index=False)
                    log_message(log_text_shipment,'文件已保存至桌面 !')
            #     SQL_Picking = f'''
            #         SELECT PLRUNNBR, PRSTYCDE, PRCOLCDE, INEXTSIZ, ACACTCDE,LOSYSLOC,SUM(BPPIDQTY)QTY
            #         FROM CTXPRWDTA.PPRPIN WHERE PLRUNNBR in({placeholders}) GROUP BY
            #         PLRUNNBR, PRSTYCDE, PRCOLCDE, INEXTSIZ, ACACTCDE,LOSYSLOC
            #         ORDER BY PLRUNNBR, PRSTYCDE, PRCOLCDE, INEXTSIZ, ACACTCDE,LOSYSLOC
            #     '''

            #     SQL_Location_Assignment = f'''
            #         SELECT A.PRSTYCDE, A.PRCOLCDE, A.INEXTSIZ, A.LOASGMIN, A.LOASGMAX,
            #         A.BPFASFLG, A.BPLOAFLG,b.losysloc FROM CTXPRWDTA.bprloa a,CTXPRWDTA.lorloc b WHERE
            #         a.lolocsgt =b.lolocsgt and a.lolocsgt not in (select lolocsgt from
            #         CTXPRWDTA.inrhld) and b.LOSYSZNE =30
            #     '''

            #     SQL_Replenishment = f'''
            #         SELECT PLRUNNBR, PRSTYCDE, PRCOLCDE, INEXTSIZ, SUM(INUPCNBR)QTY FROM CTXPRWDTA.pprsfc
            #         WHERE PLRUNNBR in({placeholders}) and ACACTCDE ='REPLEN' GROUP BY
            #         PLRUNNBR, PRSTYCDE, PRCOLCDE, INEXTSIZ, INUPCNBR ORDER BY
            #         PLRUNNBR, PRSTYCDE, PRCOLCDE, INEXTSIZ, INUPCNBR
            #     '''

            #     # 执行SQL查询
            #     data_picking = cur.execute(SQL_Picking, num_conditions).fetchall()
            #     fields_picking = [field[0] for field in cur.description]
            #     df_picking = pd.DataFrame([list(row) for row in data_picking], columns=fields_picking)

            #     if len(data_picking[0]) != len(fields_picking):
            #         raise ValueError("数据与字段名数量不匹配！")

            #     data_Location_Assignment = cur.execute(SQL_Location_Assignment).fetchall()
            #     fields_Location_Assignment = [field[0] for field in cur.description]
            #     df_Location_Assignment = pd.DataFrame([list(row) for row in data_Location_Assignment], columns=fields_Location_Assignment)

            #     # 检查数据和字段名数量是否匹配
            #     if len(data_Location_Assignment[0]) != len(fields_Location_Assignment):
            #         raise ValueError("数据与字段名数量不匹配！")
            #     data_replenishment = cur.execute(SQL_Replenishment, num_conditions).fetchall()
            #     fields_replenishment = [field[0] for field in cur.description]
            #     df_replenishment = pd.DataFrame([list(row) for row in data_replenishment], columns=fields_replenishment)

            #     if len(data_replenishment[0]) != len(fields_replenishment):
            #         raise ValueError("数据与字段名数量不匹配！")

            # # 创建一个ExcelWriter对象，使用openpyxl作为引擎
            #     output_path = os.path.join(os.path.expanduser('~'), 'Desktop', 'Run_info.xlsx')

            #     # 删除已存在的文件（如果存在）
            #     if os.path.exists(output_path):
            #         os.remove(output_path)
            #     with pd.ExcelWriter(output_path, engine='openpyxl') as writer:
            #         log_message(log_text_shipment, '文件已保存至桌面 !')
            # # 将DataFrame写入Excel的不同工作表
            #         df_Location_Assignment = process_data_with_sql(df_Location_Assignment)
            #         df_picking = process_data_with_sql(df_picking)
            #         df_replenishment = process_data_with_sql(df_replenishment)
            #         df_Location_Assignment.to_excel(writer, sheet_name='Location_Assignment', index=False)
            #         df_picking.to_excel(writer, sheet_name='Picking', index=False)
            #         df_replenishment.to_excel(writer, sheet_name='Replenishment', index=False)
                    
                    
            #     return 

            # 检查是否有6位数字
            six_digit_conditions = [cond for cond in num_conditions if len(cond) == 6]
            # 检查是否有10位数字
            ten_digit_conditions = [cond[:10] for cond in num_conditions if len(cond) > 6 and len(cond) <= 15]

            if ten_digit_conditions and not six_digit_conditions:
                # 只有10位数字，执行第一段SQL
                placeholders = ', '.join(['?'] * len(ten_digit_conditions))
                SQL = SQL_DATA['export_data']['shipment_info']
                num_conditions = ten_digit_conditions
                file_prefix = "Shipment_info"
                results = execute_sql_and_check_results(cur, SQL, num_conditions)
                if results is None:
                    SQL = SQL_DATA['export_data']['plan_info_by_por']

                    file_prefix = "plan_info"
                    results = execute_sql_and_check_results(cur, SQL, num_conditions)
                    if results is None:
                        SQL = SQL_DATA['export_data']['delay_info']
                        file_prefix = "Delay_info"

            elif six_digit_conditions and not ten_digit_conditions:
                # 只有6位数字，执行第二段SQL
                placeholders = ', '.join(['?'] * len(six_digit_conditions))
                SQL = SQL_DATA['export_data']['plan_info_by_ibrc']
                num_conditions = six_digit_conditions
                file_prefix = "plan_info"
            elif six_digit_conditions and ten_digit_conditions:
                # 既有6位数字也有10位数字，执行第三段SQL
                placeholders_6 = ', '.join(['?'] * len(six_digit_conditions))
                placeholders_10 = ', '.join(['?'] * len(ten_digit_conditions))
                SQL = SQL_DATA['export_data']['plan_info_by_both']
                num_conditions = ten_digit_conditions + six_digit_conditions
                file_prefix = "plan_info"
            if len(num_conditions[0]) == 8:
                placeholders = ', '.join(['?'] * len(num_conditions))
                SQL = SQL_DATA['export_data']['packlist_info']
                file_prefix = "Packlist_info"
            elif len(num_conditions[0]) == 20:
                placeholders = ', '.join(['?'] * len(num_conditions))
                SQL = SQL_DATA['export_data']['fs_info']
                file_prefix = "FS_info"
                results = execute_sql_and_check_results(cur, SQL, num_conditions)
                if results is None and len(num_conditions[0]) == 20 and num_conditions[0][:11] == "00000000008":
                    SQL = SQL_DATA['export_data']['trailer_info']
                    file_prefix = "trailer_info" 
            if not SQL and len(num_conditions[0]) != 4:
                log_message(log_text_shipment,'没有输入有效的查询条件。')
                return

            if len(num_conditions[0]) != 4:
                data = execute_sql_and_check_results(cur, SQL, num_conditions)
                # 获取字段名
                fields = [field[0] for field in cur.description]

                # 检查数据是否存在
                if not data:
                    log_message(log_text_shipment,'查询结果为空，没有数据可导出。')
                    return

                # 检查数据和字段名数量是否匹配
                if len(data[0]) != len(fields):
                    log_message(log_text_shipment,'数据与字段名数量不匹配，无法创建DataFrame。')
                    return
                # 创建DataFrame
                df = pd.DataFrame([list(row) for row in data], columns=fields)
                # 去除所有字符串列的前后空格和内部空格
                for col in df.select_dtypes(include='object').columns:
                    if df[col].dtype == 'O':  # 'O' 表示 object 类型，即字符串或混合类型
                        df[col] = df[col].astype(str).str.strip().str.replace(r'\s+', '', regex=True)
                        df[col] = df[col].astype(str).str.replace('None', '', regex=True)
                df = process_data_with_sql(df)

                # 获取桌面路径
                desktop_path = os.path.join(os.path.expanduser("~"), 'Desktop')
                file_name = f'{file_prefix}.xlsx'
                file_path = os.path.join(desktop_path, file_name)

                # 写入Excel文件
                df.to_excel(file_path, index=False)
                log_message(log_text_shipment,'文件已保存至桌面 !')


        if alpha_conditions:
            if len(alpha_conditions[0]) == 10:
                    SQL = SQL_DATA['export_data']['data_by_sku']
                    file_prefix = "data"
            elif  len(alpha_conditions[0]) == 9:
                    SQL = SQL_DATA['export_data']['trailer_info_by_csn']
                    file_prefix = "trailer_info"
            elif len(alpha_conditions[0]) != 10:
                    SQL = SQL_DATA['export_data']['inventory_by_sku']
                    file_prefix = "Inventory"
        # 检查SQL是否被正确赋值
        if SQL is None:
            log_message(log_text_shipment,'没有输入有效的查询条件。')
            return

        elif alpha_conditions:
            data = cur.execute(SQL, tuple(alpha_conditions)).fetchall()

            # 获取字段名
            fields = [field[0] for field in cur.description]
            # 检查数据是否存在
            if not data:
                log_message(log_text_shipment,'查询结果为空，没有数据可导出。')
                return
            # 检查数据和字段名数量是否匹配
            if len(data[0]) != len(fields):
                log_message(log_text_shipment,'数据与字段名数量不匹配，无法创建DataFrame。')
                return
            # 创建DataFrame
            df = pd.DataFrame([list(row) for row in data], columns=fields)
            # 去除所有字符串列的前后空格和内部空格
            for col in df.select_dtypes(include='object').columns:
                if df[col].dtype == 'O':  # 'O' 表示 object 类型，即字符串或混合类型
                    df[col] = df[col].astype(str).str.strip().str.replace(r'\s+', '', regex=True)
                    df[col] = df[col].astype(str).str.replace('None', '', regex=True)
            df = process_data_with_sql(df)

            # 获取桌面路径
            desktop_path = os.path.join(os.path.expanduser("~"), 'Desktop')
            file_name = f'{file_prefix}.xlsx'
            file_path = os.path.join(desktop_path, file_name)

            # 写入Excel文件
            df.to_excel(file_path, index=False)
            log_message(log_text_shipment,'文件已保存至桌面 !')

            cur.close()
            cnxn.close()

    except Exception as e:
        log_message(log_text_shipment,str(e))


# 在Notebook中显示Shipment功能
notebook.grid(row=0, column=0, sticky='w')


########
# 创建Shipping页面
shipping = ttk.Frame(notebook)

# 起止时间输入框
start_date_label = ttk.Label(shipping, text="Start Date:")
start_date_entry = ttk.Entry(shipping)
end_date_label = ttk.Label(shipping, text="End Date:")
end_date_entry = ttk.Entry(shipping)

# SQL查询类型选择
query_type = tk.StringVar(value="")
query_selector = ttk.Combobox(shipping, textvariable=query_type,
                 values=["Sorter(VNA来货)",
                         "Final-Sorter(发货)",
                         "DP & PROMO & NSRT区域补货",
                         "DP & PROMO & NSRT区域发货",
                         "异常tote",
                         "发货箱型数量(9种)",
                         "收货人效",
                         "生产信息",
                         "Cycle count data",
                         "Packlist_check",
                         "Picking",
                         "Replenishment",
                         "FM",
                         "GR",
                         "Pre_SKU_info",
                         "AP_Conveyor_info",
                         "FW_Conveyor_info",
                         "Hopper_info",
                         "Staging_info",
                         "FM_inventory",
                         "Oversize_inventory"
                         ],
                state="readonly")

# 日志输出框
log_text_shipping = tk.Text(shipping, height=10, width=50)
log_scrollbar_shipping = ttk.Scrollbar(shipping, command=log_text_shipping.yview)
log_text_shipping['yscrollcommand'] = log_scrollbar_shipping.set

queries_without_dates = [
    "AP_Conveyor_info",
    "FW_Conveyor_info",
    "Hopper_info",
    "Staging_info",
    "FM_inventory",
    "Oversize_inventory"
]
def query_and_export():
        # 连接数据库
    log_text_shipping.delete(1.0, tk.END)
    cnxn = pyodbc.connect(con_string)
    cur = cnxn.cursor()
    if cur:
        log_message(log_text_shipping,'DB 连接成功！')
    else:
        log_message(log_text_shipping,'DB 连接失败！')

    query_type_value = query_type.get()
    if query_type_value in queries_without_dates:
        start_date_entry.delete(0, tk.END)
        end_date_entry.delete(0, tk.END)
    start_date = start_date_entry.get()
    end_date = end_date_entry.get()
    try:
        if start_date and end_date:
            if query_type_value == "Sorter(VNA来货)":
                SQL = SQL_DATA['date_query']['sorter_vna']
                file_prefix = "Sorter"
            

            elif query_type_value == "Final-Sorter(发货)":
                SQL = SQL_DATA['date_query']['final_sorter']
                file_prefix = "Final-Sorter"

            elif query_type_value == "DP & PROMO & NSRT区域补货":
                SQL = SQL_DATA['date_query']['dp_promo_nsrt_replenish']
                file_prefix = "DP & PROMO & NSRT区域补货"

            elif query_type_value == "DP & PROMO & NSRT区域发货":
                SQL = SQL_DATA['date_query']['dp_promo_nsrt_shipping']
                file_prefix = "DP & PROMO & NSRT区域发货"
            elif query_type_value == "异常tote":
                SQL = SQL_DATA['date_query']['abnormal_tote']
                file_prefix = "异常tote"


            elif query_type_value == "发货箱型数量(9种)":
                SQL = SQL_DATA['date_query']['shipping_box_types']
                file_prefix = "发货箱型数量"

            elif query_type_value == "收货人效":
                SQL = SQL_DATA['date_query']['receiving_efficiency']
                file_prefix = "收货人效"

            elif query_type_value == "生产信息":
                SQL = SQL_DATA['date_query']['production_info']
                file_prefix = "生产信息"

            elif query_type_value == "Cycle count data":
                SQL = SQL_DATA['date_query']['cycle_count_data']
                file_prefix = "Cycle count data"

            elif query_type_value == "Packlist_check":
                SQL = SQL_DATA['date_query']['packlist_check']
                file_prefix = "Packlist_check"

            elif query_type_value == "Picking":
                SQL = SQL_DATA['date_query']['picking']
                file_prefix = "Picking"

            elif query_type_value == "Replenishment":
                SQL = SQL_DATA['date_query']['replenishment_qty']
                file_prefix = "Replenishment"

            elif query_type_value == "FM":
                SQL = SQL_DATA['date_query']['fm']
                file_prefix = "FM"
            
            elif query_type_value == "GR":
                SQL = SQL_DATA['date_query']['gr']
                file_prefix = "GR"
            elif query_type_value == "Pre_SKU_info":
                SQL = SQL_DATA['date_query']['pre_sku_info']
                file_prefix = "Pre_SKU_info"
        elif query_type_value == "FW_Conveyor_info":
                SQL = SQL_DATA['date_query']['fw_conveyor_info']
                file_prefix = "FW_Conveyor_info"

        elif query_type_value == "AP_Conveyor_info":
                SQL = SQL_DATA['date_query']['ap_conveyor_info']
                file_prefix = "AP_Conveyor_info"

        elif query_type_value == "Hopper_info":
                SQL = SQL_DATA['date_query']['hopper_info']
                file_prefix = "Hopper_info"

        elif query_type_value == "Staging_info":
                SQL = SQL_DATA['date_query']['staging_info']
                file_prefix = "Staging_info"

        elif query_type_value == "FM_inventory":
                SQL = SQL_DATA['date_query']['fm_inventory']
                file_prefix = "FM_inventory"

        elif query_type_value == "Oversize_inventory":
                SQL = SQL_DATA['date_query']['oversize_inventory']
                file_prefix = "Oversize_inventory"

        if start_date and end_date:
            data = cur.execute(SQL, (start_date, end_date)).fetchall()
        else:
            data = cur.execute(SQL).fetchall()
        # 获取字段名
        fields = [field[0] for field in cur.description]

        # 检查数据是否存在
        if not data:
            log_message(log_text_shipping,'查询结果为空，没有数据可导出。')
            return

        # 检查数据和字段名数量是否匹配
        if len(data[0]) != len(fields):
            log_message(log_text_shipping,'数据与字段名数量不匹配，无法创建DataFrame。')
            return

        # 创建DataFrame

        df = pd.DataFrame([list(row) for row in data], columns=fields)

        if query_type.get() == "DP & PROMO & NSRT区域补货":
            for field in fields: 
                df[field] = df[field].apply(bytes_to_string)
        elif query_type.get() == "Replenishment":
            for field in fields: 
                df[field] = df[field].apply(bytes_to_string)
        df = process_data_with_sql(df)
        # 获取桌面路径
        desktop_path = os.path.expanduser("~\\Desktop")
        file_name = f'{file_prefix}.xlsx'
        file_path = os.path.join(desktop_path, file_name)

        # 写入Excel文件
        df.to_excel(file_path, index=False)
        log_message(log_text_shipping,'文件已保存至桌面 !')
    except Exception as e:
        log_message(log_text_shipping,f'Error occurred: "请输入正确的时间区间。"')
        if cur:
            cur.close()
        if cnxn:
            cnxn.close()
def on_query_type_change(*args):
    query_type_value = query_type.get()
    if query_type_value in queries_without_dates:
        # 如果选择了一个不需要日期范围的查询类型，则清空日期输入框
        start_date_entry.delete(0, tk.END)
        end_date_entry.delete(0, tk.END)
# 导出按钮
export_button = ttk.Button(shipping, text="导出", command=query_and_export)

# 定位界面元素
start_date_label.grid(row=0, column=0, sticky='w')
start_date_entry.grid(row=0, column=1, sticky='w')
end_date_label.grid(row=1, column=0, sticky='w')
end_date_entry.grid(row=1, column=1, sticky='w')
query_selector.grid(row=2, column=0, columnspan=2, sticky='ew')
export_button.grid(row=3, column=0, columnspan=2, sticky='ew')
log_text_shipping.grid(row=4, column=0, columnspan=2, sticky='nsew')
log_scrollbar_shipping.grid(row=4, column=2, sticky='ns')
query_type.trace('w', on_query_type_change)


########


log_txt_frame = tk.Frame(master=data_export,bd=2)


log_txt = tk.Label(

    master=log_txt_frame

    # text=var_lbl,

    ,text ='Log: '

    # ,fg="white"
    # ,bg="black"

    ,width=4

    ,height=2
    )

log_txt.pack(side=tk.LEFT)


log = tk.Text(

    master=log_txt_frame
    ,fg="black"

    ,bg="white"

    ,width=100

    ,height=20
    )

log.pack(side=tk.LEFT,padx=2,pady=2)


log_scrollbar = ttk.Scrollbar(log_txt_frame, command=log.yview)

log_scrollbar.pack(side=tk.RIGHT, fill=tk.Y)


log.configure(yscrollcommand=log_scrollbar.set)

########
#
#ping_frame = tk.Frame(master=ping_ip,bd=2)
#
#
#ip_txt = tk.Label(
#    master=ping_frame
#
#    # text=var_lbl,
#
#    ,text ='请输入IP:'
#
#    # ,fg="white"
#    # ,bg="black"
#
#    ,width=8
#
#    ,anchor="w"
#    )
#
#
#IP = tk.Entry(
#    master=ping_frame
#    ,fg="black"
#
#    ,bg="white"
#
#    ,width=25
#    )
#
#
#ping = tk.Button(
#    master=ping_frame,
#
#    text="Ping",
#
#    command=Ping
#    )
#
#
#result_txt = tk.Label(
#    master=ping_frame
#
#    # text=var_lbl,
#
#    ,text ='结 果:'
#
#    # ,fg="white"
#    # ,bg="black"
#
#    ,width=8
#
#    ,anchor="w"
#    )
#
#
#Result = tk.Text(
#    master=ping_frame
#    ,fg="black"
#
#    ,bg="white"
#
#    ,width=80
#
#    ,height=20
#    )
#
#
#result_scrollbar = ttk.Scrollbar(ping_frame, command=Result.yview)
#
## result_scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
#
#result_scrollbar.grid(row=1, column=2, sticky="ns")
#
#Result.configure(yscrollcommand=result_scrollbar.set)
#
#
#
#ip_txt.grid(row=0, column=0, sticky= 'w')
#
#IP.grid(row=0, column=1, sticky= 'w')
#
#ping.grid(row=0, column=1, sticky= '', pady=5, padx= 5)
#
#result_txt.grid(row=1, column=0, sticky= 'w')
#
#Result.grid(row=1, column=1, sticky= 'w')
#
########

lable_frame = tk.Frame(master=label_history,bd=2)

query_type2 = tk.StringVar(value="")
query_selector = ttk.Combobox(lable_frame, textvariable=query_type2,
                 values=["查询打印记录","CHECK OPEN WORK","检查SKU是否是AMINUS","CHECK AUDIT WORK","QA解锁?","RSO","NFC","SO查询","查询包装信息","查询历史库位","DP_area_info","导空库位盘点数据","Export inventory"],
                 state="readonly")

# # 创建一个新的输入框用于接受特定格式的条件
# specific_condition_txt = tk.Label(
#     master=lable_frame,
#     text='请输入SKU:',
#     width=12,
#     anchor="w"
# )
# specific_input = tk.Text(
#     master=lable_frame,
#     fg="black",
#     bg="white",
#     width=80,
#     height=10
# )

carton_txt = tk.Label(
    master=lable_frame

    # text=var_lbl,

    ,text ='请输入箱号/计划号/SKU:'

    # ,fg="white"
    # ,bg="black"

    ,width=20

    ,anchor="w"
    )


carton = tk.Text(
    master=lable_frame
    ,fg="black"

    ,bg="white"

    ,width=70

    ,height=20
    )
# # 绑定事件
carton.bind("<FocusIn>", lambda event: clear_other_inputs(event, [carton]))
carton_scrollbar = ttk.Scrollbar(lable_frame, command=carton.yview)

carton_scrollbar.grid(row=0, column=2, sticky="wns")

carton.configure(yscrollcommand=carton_scrollbar.set)



query = tk.Button(
    master=lable_frame,

    text="查询",

    command=Lable_His
    )


label_log_txt = tk.Label(
    master=lable_frame

    # text=var_lbl,

    ,text ='Log:'

    # ,fg="white"
    # ,bg="black"

    ,width=9

    # ,height=2

    ,anchor="e"
    )


label_log = tk.Text(
    master=lable_frame
    ,fg="black"

    ,bg="white"

    ,width=70

    ,height=20
    )


label_log_scrollbar = ttk.Scrollbar(lable_frame, command=label_log.yview)

label_log_scrollbar.grid(row=1, column=2, sticky="wns")

label_log.configure(yscrollcommand=label_log_scrollbar.set)



# # 调整布局
# specific_condition_txt.grid(row=0, column=0, sticky='w', pady=5, padx=5)
# specific_input.grid(row=0, column=1, sticky='w', pady=5, padx=5)

# 原有的输入框布局保持不变，但是位置需要调整到新输入框下方
carton_txt.grid(row=1, column=0, sticky='w', pady=5, padx=5)
carton.grid(row=1, column=1, sticky='w', pady=5, padx=5)
query.grid(row=1, column=3, sticky='w', pady=5, padx=5)
query_selector.grid(row=1, column=4, columnspan=2, sticky='ew')
label_log_txt.grid(row=2, column=0, sticky='w', pady=5, padx=5)
label_log.grid(row=2, column=1, sticky='nw', pady=5, padx=5)


# 定义事件处理函数
def clear_other_inputs(event, other_widgets):
    for widget in other_widgets:
        widget.delete("1.0", tk.END)



########

canvas_frame = tk.Frame(master=label_history,bd=2)


##create incident 

# Add new tab for creating incidents
create_incident_frame = tk.Frame(notebook)
# API credentials and endpoint
username = SN["username"]
password = SN["password"]
urlCreate = SN["url_create"]
headers = {
    'Accept': '*/*',
    'Content-type': 'application/json'
}

# Multiple default value options
default_values = {
			"to_Node_DN":{
			"requested_by": "a.clcps", 
			"requested_for": "a.clcps", 
			"contact_type": "Self Service", 
			"requested_for_location": "1025 CLC", 
			"short_description": "1025 DN not recieved", 
			"impact": "3", 
			"urgency": "4", 
			"impacted_geo": "Greater China", 
			"cmdb_ci": "Node Management Prod",
  			"service_offering": "Node Adaptor (GC) Support",
			"assignment_group": "SCEF-GC-LOG-ADAPTOR-INC Support", 
			"detailed_description": "xx", 
			"access_key": "47642590_IIM_NKE_SN@API!2023"
			},
			"to_Node_GI":{
			"requested_by": "a.clcps", 
			"requested_for": "a.clcps", 
			"contact_type": "Self Service", 
			"requested_for_location": "1025 CLC", 
			"short_description": "1025 GI not back to SAP", 
			"impact": "3", 
			"urgency": "4", 
			"impacted_geo": "Greater China", 
			"cmdb_ci": "Node Management Prod",
  			"service_offering": "Node Adaptor (GC) Support",
			"assignment_group": "SCEF-GC-LOG-ADAPTOR-INC Support", 
			"detailed_description": "xx", 
			"access_key": "47642590_IIM_NKE_SN@API!2023"
			},
			"to_DCI":{
			"requested_by": "a.clcps", 
			"requested_for": "a.clcps", 
			"contact_type": "Self Service", 
			"requested_for_location": "1025 CLC", 
			"short_description": "1025 can't get carrier code or TN", 
			"impact": "3", 
			"urgency": "4", 
			"impacted_geo": "Greater China", 
			"cmdb_ci": "Digital Carrier Integration (DCI) (GC)",
  			"service_offering": "VCS L2 Support (GC)",
			"assignment_group": "VCS-Tech-L2-GC", 
			"detailed_description": "xx", 
			"access_key": "47642590_IIM_NKE_SN@API!2023"
			},
			"normal":{
			"requested_by": "a.clcps", 
			"requested_for": "a.clcps", 
			"contact_type": "Self Service", 
			"requested_for_location": "1025 CLC", 
			"short_description": "1025 xxxx", 
			"impact": "3", 
			"urgency": "4", 
			"impacted_geo": "Greater China", 
			"cmdb_ci": "XpDS (GC) Prod",
  			"service_offering": "WMS Legacy L1 Support (GC)",
  			"business_service": "Logistics Support (GC)",
			"assignment_group": "TechOps-LC-WMS-Legacy-L1-GC", 
			"detailed_description": "xx", 
			"access_key": "47642590_IIM_NKE_SN@API!2023"
			},
			"SAP":{
			"requested_by": "a.clcps", 
			"requested_for": "a.clcps", 
			"contact_type": "Self Service", 
			"requested_for_location": "1025 CLC", 
			"short_description": "1025 xxxx", 
			"impact": "3", 
			"urgency": "4", 
			"impacted_geo": "Greater China", 
			"cmdb_ci": "SAP S/4HANA - Order Management (OM) Prod",
  			"service_offering": "S4 Order Management L1 Support",
  			"business_service": "Nike Technology L1 Support",
			"assignment_group": "FL-S4-MSCT-Order-Management-L1L2", 
			"detailed_description": "xx", 
			"access_key": "47642590_IIM_NKE_SN@API!2023"
			}
		}

# Function to show selected default values
def show_selected_values():
    # Clear existing text
    textbox2.delete(1.0, tk.END)
    
    # Get the selected option from dropdown
    selected_option = dropdown_var.get()
    
    # Get the corresponding default values
    if selected_option in default_values:
        selected_data = default_values[selected_option]
        # Format the JSON string with double quotes
        default_json_str = json.dumps(selected_data, ensure_ascii=False, indent=4)
        textbox2.insert(tk.END, default_json_str)
    else:
        # Clear if no valid selection
        textbox2.delete(1.0, tk.END)

def create_incident(parsed_data):
    try:
        response = requests.post(urlCreate, headers=headers, auth=(username, password), json=parsed_data)
        response.raise_for_status()  # Raise an error for bad status codes
        result = response.json()
        return result
    except requests.exceptions.RequestException as e:
        return {"error": str(e)}

def on_submit():
    textbox2_value = textbox2.get("1.0", tk.END).strip()
    textbox3_value = textbox3.get("1.0", tk.END).strip()

    parsed_data = {}
    
    if textbox2_value:
        try:
            # Parse the JSON string from textbox2
            parsed_data = json.loads(textbox2_value)
            
            # If textbox3 has a value, assign it to detailed_description
            if textbox3_value:
                parsed_data["detailed_description"] = textbox3_value
        except json.JSONDecodeError:
            # Handle the case where textbox2_value is not valid JSON
            log_output.config(state=tk.NORMAL)
            log_output.delete(1.0, tk.END)
            log_output.insert(tk.END, json.dumps({
                "error": "Invalid JSON in textbox2",
                "detailed_description": textbox3_value
            }, ensure_ascii=False, indent=4))
            log_output.config(state=tk.DISABLED)
            return

    # Call the API with the constructed data
    api_response = create_incident(parsed_data)

    # Update the log output box with the API response
    log_output.config(state=tk.NORMAL)
    log_output.delete(1.0, tk.END)
    log_output.insert(tk.END, json.dumps(api_response, ensure_ascii=False, indent=4))
    log_output.config(state=tk.DISABLED)

def Show_Page(event):
    selected_tab = notebook.tab(notebook.select(), "text")
    if selected_tab == "创建incident":
        # Clear textbox2 when entering the "创建incident" tab
        textbox2.delete(1.0, tk.END)

# Assuming 'notebook' and 'create_incident_frame' are defined elsewhere in your code

# New UI elements for the "创建incident" tab
# Replace the button with a dropdown (Combobox)
dropdown_label = tk.Label(create_incident_frame, text="选择模板:")
dropdown_label.grid(row=0, column=0, sticky='w', pady=5, padx=5)

# Create a variable to store the selected dropdown value
dropdown_var = tk.StringVar()

# Create the dropdown (Combobox) with options
dropdown_options = ["请选择...", "to_Node_DN", "to_Node_GI", "to_DCI","normal","SAP"]  # Add more options as needed
dropdown = ttk.Combobox(create_incident_frame, textvariable=dropdown_var, values=dropdown_options, state="readonly")
dropdown.grid(row=0, column=1, sticky='w', pady=5, padx=5)
dropdown.current(0)  # Set default selection to "请选择..."

# Bind the selection change event
dropdown.bind("<<ComboboxSelected>>", lambda event: show_selected_values())

textbox2_label = tk.Label(create_incident_frame, text="详细描述:")
textbox2_label.grid(row=1, column=0, sticky='w', pady=5, padx=5)
textbox2 = tk.Text(create_incident_frame, height=20, width=60)
textbox2.grid(row=1, column=1, sticky='w', pady=5, padx=5)

textbox3_label = tk.Label(create_incident_frame, text="附加信息:")
textbox3_label.grid(row=2, column=0, sticky='w', pady=5, padx=5)
textbox3 = tk.Text(create_incident_frame, height=15, width=60)
textbox3.grid(row=2, column=1, sticky='w', pady=5, padx=5)

submit_button = tk.Button(create_incident_frame, text="提交", command=on_submit)
submit_button.grid(row=3, column=1, sticky='w', pady=5, padx=5)

log_output_label = tk.Label(create_incident_frame, text="API响应:")
log_output_label.grid(row=0, column=1, sticky='ne', pady=5, padx=1)
log_output = tk.Text(create_incident_frame, height=10, width=50, state=tk.DISABLED)
log_output.grid(row=1, rowspan=2, column=2, sticky='nsew', pady=5, padx=5)

# Bind the Show_Page function to the Notebook's change event
notebook.bind("<<NotebookTabChanged>>", Show_Page)

########################################################
#  window area
########################################################

option_frame.grid(row=0, column=0, sticky= 'w')

sql_frame.grid(row=1, column=0, sticky= 'w')

file_frame.grid(row=2, column=0, sticky= 'w')

folder_txt_frame.grid(row=3, column=0, sticky= 'w')

path_frame.grid(row=4, column=0, sticky= 'w')

log_txt_frame.grid(row=5, column=0, sticky= 'w')

#ping_frame.grid(row=1, column=0, sticky= 'n', pady=5, padx= 5)

lable_frame.grid(row=1, column=0, sticky= 'n', pady=5, padx= 5)

canvas_frame.grid(row=2, column=0, sticky= 'n', pady=5, padx= 5)



notebook.add(data_export, text="导数据")

#notebook.add(ping_ip, text="PING")

notebook.add(label_history, text="查历史库位/标签")

notebook.add(shipment, text="Rceiving")

notebook.add(shipping, text="按日期查询")

notebook.add(create_incident_frame, text="创建incident")

notebook.bind("<<NotebookTabChanged>>", Show_Page)

notebook.grid(row=0, column=0, sticky= 'w')   
window.mainloop()